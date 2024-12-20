import streamlit as st
import requests
import json
import time
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
import re
import random
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
import matplotlib.pyplot as plt

# Importaciones adicionales para la tabla de contenidos
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# Configuración de la página
st.set_page_config(
    page_title="Generador de Novelas de Suspenso Político",
    layout="wide",  # Uso de layout ancho para mejor visualización
    initial_sidebar_state="expanded",  # La barra lateral inicia expandida
)

# Título de la aplicación
st.title("Generador de Novelas de Suspenso Político")
st.write("""
Esta aplicación genera una novela en el género de thriller político.
Ingrese un tema y personalice el número de capítulos y escenas para crear una narrativa coherente y emocionante.
""")

# Barra lateral para opciones de usuario
st.sidebar.header("Configuración de la Novela")

# Número de capítulos
num_capitulos = st.sidebar.slider(
    "Número de capítulos",
    min_value=15,
    max_value=20,
    value=18,  # Valor recomendado
    help="Selecciona el número de capítulos para tu novela."
)

# Número de escenas por capítulo
num_escenas = st.sidebar.slider(
    "Número de escenas por capítulo",
    min_value=4,
    max_value=6,
    value=5,  # Valor recomendado
    help="Selecciona el número de escenas por cada capítulo."
)

# Porcentaje de palabras para la trama principal
porcentaje_trama_principal = st.sidebar.slider(
    "Porcentaje de palabras para la trama principal (%)",
    min_value=60,
    max_value=80,
    value=70,  # Valor recomendado
    help="Define el porcentaje de palabras que se destinarán a la trama principal."
)

# Porcentaje de palabras para subtramas
porcentaje_subtramas = 100 - porcentaje_trama_principal
st.sidebar.write(f"Porcentaje de palabras para subtramas: {porcentaje_subtramas}%")

# Inicializar el estado de la aplicación
def inicializar_estado():
    default_values = {
        'etapa': 'inicio',
        'estructura': None,
        'novela_completa': None,
        'titulo': '',
        'trama': '',
        'subtramas': '',
        'personajes': '',
        'ambientacion': '',
        'tecnica': ''
    }
    for key, value in default_values.items():
        if key not in st.session_state:
            st.session_state[key] = value

inicializar_estado()

# Función para llamar a la API de OpenRouter con reintentos y parámetros ajustables
def call_openrouter_api(prompt, max_tokens=1800, temperature=0.7, top_p=0.9, top_k=50, repetition_penalty=1.2):
    api_url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "openai/gpt-4o-mini",
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": max_tokens,
        "top_p": top_p,
        "top_k": top_k,
        "repetition_penalty": repetition_penalty,
        "stop": ["[\"<|eot_id|>\"]"],
        "stream": False
    }
    
    session = requests.Session()
    retries = Retry(total=5, backoff_factor=1, status_forcelist=[502, 503, 504])
    session.mount('https://', HTTPAdapter(max_retries=retries))
    
    try:
        response = session.post(api_url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
        response_json = response.json()
        if 'choices' in response_json and len(response_json['choices']) > 0:
            return response_json['choices'][0]['message']['content']
        else:
            st.error("La respuesta de la API no contiene 'choices'.")
            st.write("Respuesta completa de la API:", response_json)
            return None
    except requests.exceptions.RequestException as e:
        st.error(f"Error en la llamada a la API: {e}")
        return None

# Función para generar la estructura inicial de la novela con subtramas y técnicas avanzadas
def generar_estructura(theme):
    prompt = f"""
Basado en el tema proporcionado, genera una estructura detallada para una novela de suspenso político de alta calidad. Asegúrate de que la novela obtenga una calificación de 10 sobre 10 en los siguientes aspectos:
- **Trama**: Compleja, bien desarrollada y llena de giros inesperados.
- **Originalidad**: Ideas frescas y únicas que distinguen la novela de otras en el mismo género.
- **Desarrollo de Personajes**: Personajes profundos, multidimensionales y realistas con arcos de desarrollo claros.
- **Ritmo**: Fluido y bien equilibrado, manteniendo el interés del lector en todo momento.
- **Descripciones**: Vivas y detalladas que permiten al lector visualizar escenas y emociones con claridad, sin extenderse demasiado. **Evita frases hechas** como "un silencio ensordecedor" o "el corazón latía apresuradamente".
- **Calidad General**: Cohesión, coherencia y excelencia literaria en todo momento.
- **Técnicas Avanzadas de Escritura**:
    - **Foreshadowing**: Introduce pistas sutiles sobre eventos futuros.
    - **Metáforas y Simbolismo**: Utiliza figuras retóricas para enriquecer la narrativa.
    - **Show, Don't Tell**: Enfócate en mostrar acciones y emociones en lugar de simplemente describirlas.

### Estructura Requerida:
1. **Título**
2. **Trama Principal**
3. **Subtramas** (incluyendo nombres, descripciones detalladas, motivaciones y cómo afectan a los personajes y la trama principal)
4. **Personajes** (incluyendo nombres, descripciones físicas y psicológicas, motivaciones, y arcos de desarrollo)
5. **Ambientación** (detallada y relevante para la trama)
6. **Técnicas Literarias a Utilizar** (como metáforas, simbolismo, foreshadowing, etc.)

### Tema:
{theme}

Asegúrate de que toda la información generada sea coherente y adecuada para un thriller político de alta calidad.
"""
    estructura = call_openrouter_api(prompt)
    return estructura

# Función para extraer los elementos de la estructura usando expresiones regulares
def extraer_elementos(estructura):
    # Mostrar la estructura completa para depuración
    st.write("### Estructura generada por la API:")
    st.write(estructura)

    # Patrón mejorado para extraer los elementos, incluyendo subtramas
    patrones = {
        'titulo': r"(?:Título|Titulo):\s*(.*)",
        'trama': r"Trama Principal:\s*((?:.|\n)*?)\n(?:Subtramas|Subtrama)",
        'subtramas': r"Subtramas?:\s*((?:.|\n)*?)\n(?:Personajes|Ambientación|Ambientacion|Técnicas literarias|$)",
        'personajes': r"Personajes:\s*((?:.|\n)*?)\n(?:Ambientación|Ambientacion|Técnicas literarias|$)",
        'ambientacion': r"Ambientación:\s*((?:.|\n)*?)\n(?:Técnicas literarias|$)",
        'tecnica': r"Técnicas literarias(?: a utilizar)?:\s*((?:.|\n)*)"
    }

    titulo = re.search(patrones['titulo'], estructura, re.IGNORECASE)
    trama = re.search(patrones['trama'], estructura, re.IGNORECASE)
    subtramas = re.search(patrones['subtramas'], estructura, re.IGNORECASE)
    personajes = re.search(patrones['personajes'], estructura, re.IGNORECASE)
    ambientacion = re.search(patrones['ambientacion'], estructura, re.IGNORECASE)
    tecnica = re.search(patrones['tecnica'], estructura, re.IGNORECASE)

    # Extraer el contenido, manejando posibles espacios y formatos
    titulo = titulo.group(1).strip() if titulo else "Sin título"
    trama = trama.group(1).strip() if trama else "Sin trama principal"
    subtramas = subtramas.group(1).strip() if subtramas else "Sin subtramas"
    personajes = personajes.group(1).strip() if personajes else "Sin personajes"
    ambientacion = ambientacion.group(1).strip() if ambientacion else "Sin ambientación"
    tecnica = tecnica.group(1).strip() if tecnica else "Sin técnicas literarias"

    return titulo, trama, subtramas, personajes, ambientacion, tecnica

# Función para generar cada escena con subtramas y técnicas avanzadas de escritura
def generar_escena(capitulo, escena, trama, subtramas, personajes, ambientacion, tecnica, palabras_trama, palabras_subtramas):
    # Estimar tokens: 1 palabra ≈ 1.3 tokens
    max_tokens_trama = int(palabras_trama * 1.3)
    max_tokens_subtramas = int(palabras_subtramas * 1.3)
    total_max_tokens = max_tokens_trama + max_tokens_subtramas

    prompt = f"""
Escribe la Escena {escena} del Capítulo {capitulo} de una novela de suspenso político de alta calidad con las siguientes características:

- **Trama Principal**: {trama}
- **Subtramas**: {subtramas}
- **Personajes**: {personajes}
- **Ambientación**: {ambientacion}
- **Técnicas Literarias**: {tecnica}

### Requisitos de la Escena:
1. **Trama**: Desarrolla la trama principal con profundidad y añade giros inesperados que mantengan al lector intrigado.
2. **Subtramas**: Integra las subtramas de manera que complementen y enriquezcan la trama principal, asegurando que cada una contribuya al desarrollo de los personajes y al avance de la historia.
3. **Desarrollo de Personajes**: Asegúrate de que las interacciones entre personajes muestren sus arcos de desarrollo y relaciones complejas.
4. **Ritmo**: Mantén un ritmo dinámico que equilibre la acción, el suspense y el desarrollo emocional.
5. **Descripciones**: Utiliza descripciones vivas y detalladas, evitando que sean demasiado extensas. **Evita frases hechas y comunes** como “un silencio ensordecedor” o “el corazón latía apresuradamente”. **No repitas frases o clichés**; procura que cada descripción aporte frescura y claridad sin extenderse demasiado.
6. **Inicio de Escena**: Los inicios de escena deben ser originales y no predecibles, evitando que se repitan patrones o comienzos que resulten similares en escenas consecutivas.
7. **Detalles de Personajes**: Incluye detalles sutiles de la vida pasada de los personajes, pensamientos internos o conflictos personales que expliquen sus acciones y decisiones. Asegúrate de que cada personaje tenga una voz única, usando pequeñas descripciones o acciones para dar contexto a sus palabras y reflejar su personalidad.
8. **Calidad Literaria**: Emplea técnicas literarias avanzadas como metáforas, simbolismo y foreshadowing para enriquecer la narrativa.
9. **Coherencia y Cohesión**: Asegúrate de que los eventos y desarrollos sean lógicos y estén bien conectados con el resto de la historia.
10. **Condensación de Escenas**: Condensa las escenas que no añadan mucha información nueva y enfócate en momentos clave que impulsen la historia o revelen aspectos críticos de la trama.
11. **Vínculos con Giros**: Vincula cada escena de tensión o suspenso con un giro importante o una revelación para que el lector perciba un avance constante en la historia.
12. **Variedad de Lenguaje**: Usa sinónimos o reformula las ideas para evitar la repetición y hacer el texto más variado.
13. **Optimización de Descripciones**: Identifica las descripciones esenciales para el ambiente y reduce o simplifica las que no aportan directamente al desarrollo de la trama.
14. **Orientación del Lector**: Añadir nombres de personajes o pequeñas indicaciones al cambiar de escena o personaje para mantener al lector orientado.
15. **Consistencia de Tono**: Mantén un tono consistente, especialmente en escenas que comparten una misma intensidad emocional. Si se va a cambiar, utiliza un recurso narrativo claro para evitar confusión.
16. **Fluidez entre Escenas**: Asegúrate de que cada escena fluya suavemente a la siguiente. Podrías usar frases que conecten o introduzcan el cambio en el espacio o la acción para guiar mejor al lector.
17. **Naturalidad de Acciones**: Asegúrate de que las acciones de los personajes se sientan naturales. Proporciona pistas o justificaciones previas para sus movimientos estratégicos.

### Distribución de Palabras:
- **Trama Principal**: Aproximadamente {palabras_trama} palabras.
- **Subtramas**: Aproximadamente {palabras_subtramas} palabras.

### Formato:
- Utiliza rayas (—) para las intervenciones de los personajes.
- Estructura el texto con párrafos claros y bien organizados.
- Evita clichés y frases hechas, enfocándote en originalidad y frescura.

Asegúrate de mantener la coherencia y la cohesión en toda la escena, contribuyendo significativamente al desarrollo general de la novela.
"""
    escena_texto = call_openrouter_api(prompt, max_tokens=total_max_tokens, temperature=0.7, top_p=0.9, top_k=50, repetition_penalty=1.2)
    return escena_texto

# Función para generar la novela completa después de la aprobación
def generar_novela_completa(num_capitulos, num_escenas):
    titulo = st.session_state.titulo
    trama = st.session_state.trama
    subtramas = st.session_state.subtramas
    personajes = st.session_state.personajes
    ambientacion = st.session_state.ambientacion
    tecnica = st.session_state.tecnica

    total_palabras = 60000  # Ajustado a 60,000 palabras
    total_escenas = num_capitulos * num_escenas

    # Distribuir las palabras entre trama principal y subtramas
    porcentaje_trama_principal_decimal = porcentaje_trama_principal / 100  # Convertir a decimal
    porcentaje_subtramas_decimal = porcentaje_subtramas / 100

    palabras_trama_principal_total = int(total_palabras * porcentaje_trama_principal_decimal)
    palabras_subtramas_total = total_palabras - palabras_trama_principal_total

    # Distribuir palabras por escena para trama principal
    palabras_por_escena_trama = palabras_trama_principal_total // total_escenas
    palabras_restantes_trama = palabras_trama_principal_total - (palabras_por_escena_trama * total_escenas)

    # Distribuir palabras por escena para subtramas
    palabras_por_escena_subtramas = palabras_subtramas_total // total_escenas
    palabras_restantes_subtramas = palabras_subtramas_total - (palabras_por_escena_subtramas * total_escenas)

    # Crear listas de palabras por escena con variación del ±75 palabras para trama y ±45 para subtramas
    palabras_por_escena_trama_lista = []
    palabras_por_escena_subtramas_lista = []
    for _ in range(total_escenas):
        variacion_trama = random.randint(-75, 75)  # Variación recomendada
        palabras_trama = palabras_por_escena_trama + variacion_trama
        palabras_trama = max(350, palabras_trama)  # Mínimo recomendado
        palabras_por_escena_trama_lista.append(palabras_trama)

        variacion_subtramas = random.randint(-45, 45)  # Variación recomendada
        palabras_subtramas = palabras_por_escena_subtramas + variacion_subtramas
        palabras_subtramas = max(175, palabras_subtramas)  # Mínimo recomendado
        palabras_por_escena_subtramas_lista.append(palabras_subtramas)

    # Ajustar las palabras restantes
    for i in range(palabras_restantes_trama):
        palabras_por_escena_trama_lista[i % total_escenas] += 1

    for i in range(palabras_restantes_subtramas):
        palabras_por_escena_subtramas_lista[i % total_escenas] += 1

    novela = f"**{titulo}**\n\n"

    # Inicializar la barra de progreso
    progress_bar = st.progress(0)
    progress_text = st.empty()
    current = 0

    escena_index = 0  # Índice para acceder a las listas de palabras
    palabras_por_capitulo = {cap: [] for cap in range(1, num_capitulos + 1)}  # Para la gráfica

    # Generar cada capítulo y escena
    for cap in range(1, num_capitulos + 1):
        novela += f"## Capítulo {cap}\n\n"
        for esc in range(1, num_escenas + 1):
            palabras_trama_escena = palabras_por_escena_trama_lista[escena_index]
            palabras_subtramas_escena = palabras_por_escena_subtramas_lista[escena_index]
            total_palabras_escena = palabras_trama_escena + palabras_subtramas_escena

            palabras_por_capitulo[cap].append(total_palabras_escena)
            with st.spinner(f"Generando Capítulo {cap}, Escena {esc} ({total_palabras_escena} palabras)..."):
                escena = generar_escena(cap, esc, trama, subtramas, personajes, ambientacion, tecnica, 
                                        palabras_trama_escena, palabras_subtramas_escena)
                if not escena:
                    st.error(f"No se pudo generar la Escena {esc} del Capítulo {cap}.")
                    return None
                # Limpiar saltos de línea manuales, reemplazándolos por saltos de párrafo
                escena = escena.replace('\r\n', '\n').replace('\n', '\n\n')
                novela += f"### Escena {esc}\n\n{escena}\n\n"
                # Actualizar la barra de progreso
                current += 1
                progress_bar.progress(current / total_escenas)
                progress_text.text(f"Progreso: {current}/{total_escenas} escenas generadas.")
                escena_index += 1
                # Retraso para evitar exceder los límites de la API
                time.sleep(1)

    # Ocultar la barra de progreso y el texto de progreso
    progress_bar.empty()
    progress_text.empty()

    # Mostrar el total de palabras generadas estimadas
    total_palabras_generadas = len(novela.split())
    st.write(f"**Total de palabras generadas estimadas:** {total_palabras_generadas}")

    # Graficar la distribución de palabras por capítulo
    fig, ax = plt.subplots(figsize=(10, 6))
    for cap in palabras_por_capitulo:
        ax.plot(
            range(1, num_escenas + 1),
            palabras_por_capitulo[cap],
            marker='o',
            label=f'Capítulo {cap}'
        )
    ax.set_xlabel('Escena')
    ax.set_ylabel('Palabras')
    ax.set_title('Distribución de Palabras por Escena en Cada Capítulo')
    ax.legend()
    st.pyplot(fig)

    st.session_state.novela_completa = novela
    st.session_state.etapa = "completado"

    return novela

# Función para exportar la novela a un archivo de Word con el formato especificado
def exportar_a_word(titulo, novela_completa):
    document = Document()

    # Configurar el tamaño de la página y márgenes
    section = document.sections[0]
    section.page_width = Inches(5)
    section.page_height = Inches(8)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    # Establecer el estilo normal con una fuente común
    style = document.styles['Normal']
    font = style.font
    font.name = 'Alegreya'  # Fuente recomendada para documentos formales
    font.size = Pt(11)

    # Agregar el título
    titulo_paragraph = document.add_heading(titulo, level=0)
    titulo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Agregar la tabla de contenidos
    agregar_tabla_de_contenidos(document)
    document.add_page_break()

    # Separar la novela por capítulos
    capítulos = novela_completa.split("## Capítulo")
    for cap in capítulos:
        cap = cap.strip()
        if not cap:
            continue
        cap_num_match = re.match(r"(\d+)", cap)
        cap_num = cap_num_match.group(1) if cap_num_match else "Sin número"
        cap_content = cap.split('\n', 1)[1].strip() if '\n' in cap else ""

        # Agregar el capítulo
        document.add_heading(f"Capítulo {cap_num}", level=1)

        # Separar por escenas
        escenas = cap_content.split("### Escena")
        for esc in escenas:
            esc = esc.strip()
            if not esc:
                continue
            esc_num_match = re.match(r"(\d+)", esc)
            esc_num = esc_num_match.group(1) if esc_num_match else "Sin número"
            esc_text = esc.split('\n', 1)[1].strip() if '\n' in esc else ""

            # Agregar la escena
            document.add_heading(f"Escena {esc_num}", level=2)

            # Agregar el texto de la escena con saltos de párrafo
            for paragraph_text in esc_text.split('\n\n'):
                paragraph_text = paragraph_text.strip()
                if paragraph_text:
                    paragraph = document.add_paragraph(paragraph_text)
                    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
                    paragraph_format = paragraph.paragraph_format
                    paragraph_format.line_spacing = 1.15
                    paragraph_format.space_after = Pt(6)

    # Agregar el conteo total de palabras al final del documento
    total_palabras = len(novela_completa.split())
    document.add_page_break()
    document.add_paragraph(f"**Total de palabras generadas:** {total_palabras}", style='Intense Quote')

    # Guardar el documento en memoria
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Función para agregar una tabla de contenidos automática
def agregar_tabla_de_contenidos(document):
    paragraph = document.add_paragraph()
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

# Interfaz de usuario para aprobar la estructura inicial
def mostrar_aprobacion():
    st.header("Aprobación de Elementos Iniciales")
    st.subheader("Título")
    st.write(st.session_state.titulo)

    st.subheader("Trama Principal")
    st.write(st.session_state.trama)

    st.subheader("Subtramas")
    st.write(st.session_state.subtramas)

    st.subheader("Personajes")
    st.write(st.session_state.personajes)

    st.subheader("Ambientación")
    st.write(st.session_state.ambientacion)

    st.subheader("Técnicas Literarias")
    st.write(st.session_state.tecnica)

    # Botones de aprobación y rechazo
    aprobar, rechazar = st.columns(2)
    with aprobar:
        if st.button("Aprobar y Generar Novela", key="aprobar"):
            st.session_state.etapa = "generacion"
    with rechazar:
        if st.button("Rechazar y Regenerar Estructura", key="rechazar"):
            # Reiniciamos los valores
            st.session_state.estructura = None
            st.session_state.titulo = ""
            st.session_state.trama = ""
            st.session_state.subtramas = ""
            st.session_state.personajes = ""
            st.session_state.ambientacion = ""
            st.session_state.tecnica = ""
            st.session_state.etapa = "inicio"

# Interfaz de usuario principal
st.write(f"**Etapa actual:** {st.session_state.etapa}")  # Información de depuración

if st.session_state.etapa == "inicio":
    st.header("Generación de Elementos Iniciales")
    theme = st.text_input("Ingrese el tema para su thriller político:", "")

    if st.button("Generar Elementos Iniciales"):
        if not theme:
            st.error("Por favor, ingrese un tema.")
        else:
            with st.spinner("Generando la estructura inicial..."):
                estructura = generar_estructura(theme)
                if estructura:
                    titulo, trama, subtramas, personajes, ambientacion, tecnica = extraer_elementos(estructura)
                    # Guardar en el estado de la sesión
                    st.session_state.estructura = estructura
                    st.session_state.titulo = titulo
                    st.session_state.trama = trama
                    st.session_state.subtramas = subtramas
                    st.session_state.personajes = personajes
                    st.session_state.ambientacion = ambientacion
                    st.session_state.tecnica = tecnica
                    st.session_state.etapa = "aprobacion"
                else:
                    st.error("No se pudo generar la estructura inicial. Por favor, intente nuevamente.")

if st.session_state.etapa == "aprobacion":
    mostrar_aprobacion()

if st.session_state.etapa == "generacion":
    with st.spinner("Generando la novela completa..."):
        novela_completa = generar_novela_completa(num_capitulos, num_escenas)
        if novela_completa:
            st.session_state.etapa = "completado"

if st.session_state.etapa == "completado":
    if st.session_state.novela_completa:
        st.success("Novela generada con éxito.")
        # Exportar a Word
        doc_buffer = exportar_a_word(st.session_state.titulo, st.session_state.novela_completa)
        st.download_button(
            label="Descargar Novela en Word",
            data=doc_buffer,
            file_name=f"novela_thriller_politico_{int(time.time())}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        # Mostrar la novela en la interfaz
        st.text_area("Novela Generada:", st.session_state.novela_completa, height=600)
        st.info("Después de descargar el documento en Word, abre el archivo y actualiza la tabla de contenidos:")
        st.markdown("""
1. Selecciona la tabla de contenidos.
2. Haz clic derecho y selecciona "Actualizar campo".
3. Elige "Actualizar toda la tabla" y confirma.
        """)

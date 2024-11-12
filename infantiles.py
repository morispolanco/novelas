import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO
import random
from PIL import Image
import base64
import nltk
from nltk.tokenize import sent_tokenize
import concurrent.futures

# Configuración de la página
st.set_page_config(
    page_title="Generador de Cuentos Infantiles con Ilustraciones",
    layout="centered",
    initial_sidebar_state="auto",
)

# Función para descargar recursos de NLTK
@st.cache_resource
def download_nltk_resources():
    try:
        nltk.data.find('tokenizers/punkt')
    except LookupError:
        nltk.download('punkt', quiet=True)

# Descargar recursos al iniciar la aplicación
download_nltk_resources()

# Título de la aplicación
st.title("📚 Generador de Cuentos Infantiles con Ilustraciones")

# Descripción
st.markdown("""
Genera cuentos personalizados en español para niños, adaptados a diferentes grupos de edad. 
Especifica el número de cuentos (hasta 15), y deja que la aplicación cree historias atractivas con temas variados, longitudes apropiadas y dos hermosas ilustraciones por cuento.
""")

# Lista predefinida de 150 temas
themes = [
    # [Lista completa de temas]
    "aventura en el bosque",
    "un viaje mágico",
    "amistad y trabajo en equipo",
    # ... (resto de los temas)
    "el mago del bosque"
]

# Lista predefinida de nombres de personajes únicos
character_names = [
    "Luna", "Mateo", "Sofía", "Emilio", "Valentina", "Lucas", "Isabella", "Alejandro",
    "Camila", "Gabriel", "Emma", "Benjamín", "Victoria", "Daniel", "Lucía", "Diego",
    "Martina", "Samuel", "Natalia", "Sebastián", "Valeria", "Emiliano", "Catalina",
    "Amelia", "Jorge", "Renata", "Andrés", "Sara", "Antonio", "Claudia",
    "Pablo", "Mía", "Ricardo", "Alicia", "Javier", "Paula", "Santiago", "Gabriela",
    "Hugo", "María", "Fernando", "Julia", "Adrián", "Lorena", "Tomás", "Andrea",
    "Óscar", "Fernanda"
]

# Función para generar un único nombre de personaje
def get_unique_name(used_names):
    available_names = list(set(character_names) - set(used_names))
    if not available_names:
        # Si se agotan los nombres, generar uno aleatorio
        return f"Personaje_{random.randint(1000, 9999)}"
    return random.choice(available_names)

# Función para extraer momentos clave de la historia
def extract_key_moments(story, num_moments=2):
    sentences = sent_tokenize(story, language='spanish')
    if len(sentences) < num_moments:
        return sentences
    interval = len(sentences) // (num_moments + 1)
    moments = [sentences[interval * (i + 1)] for i in range(num_moments)]
    return moments

# Barra lateral para la entrada del usuario
st.sidebar.header("Parámetros de Entrada")

# Entrada para el número de cuentos
num_stories = st.sidebar.number_input(
    "Número de Cuentos",
    min_value=1,
    max_value=15,
    value=1,
    step=1
)

# Grupos de edad para elegir
age_groups = ["3-5 años", "6-8 años", "9-12 años"]
selected_age_group = st.sidebar.selectbox(
    "Selecciona el Grupo de Edad para los Cuentos",
    age_groups
)

# Botón para generar cuentos
if st.sidebar.button("Generar Cuentos"):
    with st.spinner("Generando cuentos e ilustraciones..."):
        # Función para generar y corregir un solo cuento usando OpenRouter
        def generate_corrected_story(theme, age_group, character_name):
            api_url = "https://openrouter.ai/api/v1/chat/completions"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
            }

            # Definir las 20 directrices
            guidelines = """
1. **Desarrollo de Personajes:**
   - Profundiza en los conflictos internos de los personajes. Muestra sus miedos, dudas y pequeñas victorias para hacerlos más humanos y relacionables. A medida que cada personaje enfrenta un desafío, incluye momentos donde reflexionen sobre sus inseguridades o limitaciones para enriquecer su crecimiento.

2. **Consistencia en el Tono y Edad:**
   - Las historias están dirigidas a lectores de 12 años, pero algunos temas y emociones pueden ser ligeramente avanzados para esta edad. Simplifica ciertas decisiones morales y emocionales, o haz que los personajes actúen y piensen de manera apropiada para su edad para mantener la autenticidad.

3. **Detalles Sensoriales:**
   - Utiliza descripciones sensoriales para hacer que los entornos y experiencias sean más inmersivos. Incorpora olores, sonidos y texturas para crear una atmósfera más rica y ayudar a los lectores a sentirse más conectados con el mundo de la historia.

4. **Morales Sutiles:**
   - Presenta las lecciones de cada historia de manera más implícita en lugar de explícita. Permite que los personajes y los lectores descubran la moral a través de acciones y consecuencias, haciendo que el mensaje sea más efectivo y menos moralizante.

5. **Profundización de Relaciones entre Personajes:**
   - Los lazos entre personajes añaden una capa emocional a las historias. Muestra interacciones que revelen más sobre sus relaciones (conflictos, apoyo, amistades) para ayudar a los lectores a sentirse más conectados emocionalmente con la historia.

6. **Conflictos y Desafíos Graduales:**
   - Asegúrate de que los personajes enfrenten obstáculos en una progresión natural. Esto añade tensión y hace que la resolución sea más satisfactoria, ya que el personaje supera diversos desafíos para lograr su objetivo.

7. **Variedad de Conflictos:**
   - Varía los tipos de conflictos (internos, externos, naturales, sobrenaturales, emocionales) para dar dinamismo a la colección y evitar un patrón repetitivo. Esto permite la exploración de diferentes tipos de coraje, empatía y creatividad.

8. **Símbolos y Elementos Recurrentes:**
   - Incluye símbolos o elementos recurrentes (como objetos, frases o lugares) que tengan un significado especial en cada historia. Esto unifica la colección temáticamente, ayudando a los lectores a encontrar conexiones y significados a lo largo de las historias.

9. **Contexto del Mundo y Antecedentes:**
   - En historias con un mundo mágico o fantástico, ofrece un breve contexto del entorno o las "reglas mágicas" para construir un universo consistente. Esto mejora el atractivo y la inmersión de la historia.

10. **Mostrar, No Contar:**
    - En lugar de decir directamente al lector lo que un personaje siente o aprende, muéstralo a través de acciones, gestos y decisiones. Esta técnica narrativa permite que los lectores interpreten las emociones del personaje y extraigan la moral por sí mismos.

11. **Fortalecer Temas de Fondo:**
    - Asegúrate de que cada historia tenga un tema central claro (como valentía, amistad o empatía) y explóralo en profundidad. Los temas consistentes a lo largo de la colección le dan una identidad cohesiva.

12. **Desarrollar Personajes Memorables:**
    - Los personajes principales deben tener cualidades distintivas (como valentía, ingenio o humildad) que se reflejen en sus acciones y decisiones. Añade pequeños detalles físicos o emocionales para hacerlos más humanos y fáciles de recordar para los lectores.

13. **Agregar Detalles Sensoriales:**
    - Utiliza descripciones sensoriales (olor, tacto, sonido) para sumergir a los lectores en el entorno de cada historia. Estos detalles enriquecen la narrativa y ayudan a los lectores a sentir que están explorando estos mundos con los personajes.

14. **Incorporar Conflictos o Desafíos Claros:**
    - Cada historia debe tener un conflicto o desafío central para que el protagonista lo supere. Esto crea tensión y mantiene a los lectores enganchados. Además, resolver problemas o enfrentar dilemas hace que las historias sean más dinámicas.

15. **Mostrar Crecimiento y Aprendizaje:**
    - Asegúrate de que los personajes crezcan o aprendan algo significativo en cada historia. Las lecciones o cambios internos deben sentirse naturales y reflejarse en la narrativa sin parecer forzados, haciendo que el aprendizaje sea algo que los lectores puedan sentir y comprender.

16. **Crear una Conexión con el Lector:**
    - Utiliza situaciones o emociones que los niños puedan relacionar, como la curiosidad, el miedo a lo desconocido o la búsqueda de aceptación. Esto les ayuda a sentirse parte de la historia y a comprender mejor el mensaje.

17. **Añadir Humor y Toques Ligeros:**
    - Aunque las historias deben contener una lección, añade humor en diálogos o interacciones entre personajes. Esto no solo hace que la lectura sea más entretenida, sino que también equilibra la historia, especialmente si el mensaje es profundo.

18. **Usar Ritmo y Variedad en la Narrativa:**
    - Juega con el ritmo, alternando momentos de calma con picos de emoción o tensión. Además, varía el lenguaje y el estilo de narración en cada historia para mantener la frescura, incluso con un tono general consistente.

19. **Crear Introducciones y Conclusiones Memorables:**
    - Las líneas de apertura deben captar la atención del lector, y las conclusiones deben dejar una sensación de satisfacción o reflexión. Esto hace que la historia sea más impactante y memorable.

20. **Transmitir Valores Positivos sin Ser Moralista:**
    - En lugar de hacer la moral explícita, permite que los valores positivos emerjan naturalmente de las decisiones y acciones de los personajes. Esto evita que las historias se sientan moralizantes y da al mensaje una resonancia poderosa y sutil.
            """

            # Añadir las directrices al prompt
            prompt = f"""
{guidelines}

Crea una historia que cumpla con las siguientes características:

- Tema: {theme}
- Grupo de edad: {age_group}
- Nombre del personaje principal: {character_name}
- Texto continuo sin subdivisiones ni subtítulos.
            """

            data = {
                "model": "gpt-4",  # Asegúrate de que este modelo esté disponible en OpenRouter
                "messages": [
                    {
                        "role": "user",
                        "content": prompt
                    }
                ],
                "temperature": 0.7,
                "max_tokens": 1000  # Ajusta según sea necesario
            }

            try:
                response = requests.post(api_url, headers=headers, data=json.dumps(data))
                response.raise_for_status()
                result = response.json()

                # Verificar estructura de la respuesta
                if 'choices' in result and len(result['choices']) > 0:
                    story = result['choices'][0]['message']['content'].strip()
                else:
                    st.error("Respuesta inesperada de la API de OpenRouter.")
                    return "Lo siento, ocurrió un error al generar el cuento."

                return story
            except requests.exceptions.HTTPError as http_err:
                st.error(f"Ocurrió un error HTTP: {http_err}")
                if 'response' in locals():
                    st.write(response.text)  # Mostrar la respuesta completa para depuración
            except Exception as err:
                st.error(f"Ocurrió un error: {err}")
            return "Lo siento, ocurrió un error al generar el cuento."

        # Función para generar una ilustración usando Together.xyz
        def generate_image(prompt_description):
            together_api_key = st.secrets.get('TOGETHER_API_KEY')
            if not together_api_key:
                st.error("La clave API de Together.xyz no está configurada en los secretos de Streamlit.")
                return None

            api_url = "https://api.together.xyz/v1/images/generations"
            headers = {
                "Content-Type": "application/json",
                "Authorization": f"Bearer {together_api_key}"
            }

            data = {
                "model": "black-forest-labs/FLUX.1.1-pro",
                "prompt": prompt_description,
                "width": 512,
                "height": 512,
                "steps": 1,
                "n": 1,
                "response_format": "b64_json"
            }

            try:
                response = requests.post(api_url, headers=headers, data=json.dumps(data))
                response.raise_for_status()
                result = response.json()
                b64_image = result['data'][0]['b64_json']
                image_bytes = base64.b64decode(b64_image)
                image = Image.open(BytesIO(image_bytes))
                return image
            except requests.exceptions.HTTPError as http_err:
                st.error(f"Ocurrió un error HTTP al generar la imagen: {http_err}")
                if 'response' in locals():
                    st.write(response.text)  # Mostrar la respuesta completa para depuración
            except Exception as err:
                st.error(f"Ocurrió un error al generar la imagen: {err}")
            return None

        # Seleccionar temas para el número de cuentos
        selected_themes = random.sample(themes, num_stories) if num_stories <= len(themes) else random.choices(themes, k=num_stories)

        # Generar todos los cuentos y sus ilustraciones
        stories = []
        used_names = []
        progress_bar = st.progress(0)

        # Uso de ThreadPoolExecutor para paralelizar la generación de cuentos e imágenes
        with concurrent.futures.ThreadPoolExecutor() as executor:
            future_to_story = {}
            for i in range(num_stories):
                theme = selected_themes[i]
                # Obtener un nombre de personaje único
                character_name = get_unique_name(used_names)
                used_names.append(character_name)
                # Iniciar la generación del cuento
                future = executor.submit(generate_corrected_story, theme, selected_age_group, character_name)
                future_to_story[future] = (theme, character_name, i)

            for future in concurrent.futures.as_completed(future_to_story):
                theme, character_name, i = future_to_story[future]
                story = future.result()
                if story.startswith("Lo siento"):
                    images = [None, None]
                else:
                    # Extraer momentos clave de la historia
                    key_moments = extract_key_moments(story)
                    if len(key_moments) < 2:
                        # Si hay menos de 2 momentos clave, usar descripciones generales
                        image_prompt1 = f"Ilustración colorida y atractiva para un cuento infantil sobre {theme}."
                        image_prompt2 = f"Ilustración que represente un momento significativo en la historia de {character_name} en el tema de {theme}."
                    else:
                        image_prompt1 = f"Ilustración colorida y atractiva de un momento clave: {key_moments[0]}"
                        image_prompt2 = f"Ilustración que represente otro momento significativo: {key_moments[1]}"
                    # Generar las imágenes
                    image1 = generate_image(image_prompt1)
                    image2 = generate_image(image_prompt2)
                    images = [image1, image2]
                stories.append({
                    "title": f"Capítulo {i+1}: {theme.title()}",
                    "content": story,
                    "images": images  # Almacenar una lista de imágenes
                })
                progress_bar.progress((i + 1) / num_stories)

        # Crear un documento de Word
        doc = Document()
        doc.add_heading("Cuentos Infantiles con Ilustraciones", 0)

        for story in stories:
            doc.add_heading(story["title"], level=1)
            doc.add_paragraph(story["content"])
            if story["images"]:
                for idx, image in enumerate(story["images"]):
                    if image:
                        # Guardar la imagen en un flujo de bytes
                        img_byte_arr = BytesIO()
                        image.save(img_byte_arr, format='PNG')
                        img_byte_arr = img_byte_arr.getvalue()
                        
                        # Agregar la imagen al documento
                        doc.add_picture(BytesIO(img_byte_arr), width=Inches(4))
                        doc.add_paragraph("")  # Espacio adicional después de la imagen
                    else:
                        doc.add_paragraph(f"No se pudo generar la ilustración {idx + 1}.")
            else:
                doc.add_paragraph("No se pudieron generar las ilustraciones para este cuento.")

        # Guardar el documento en un flujo de BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

        # Proporcionar el botón de descarga
        st.success("¡Cuentos e ilustraciones generados con éxito!")
        st.download_button(
            label="📄 Descargar Cuentos como Documento de Word",
            data=doc_io,
            file_name="Cuentos_Infantiles_Con_Ilustraciones.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

        # Opcional: mostrar los cuentos e ilustraciones en la aplicación
        with st.expander("Ver Cuentos e Ilustraciones Generados"):
            for story in stories:
                st.markdown(f"### {story['title']}")
                st.write(story["content"])
                if story["images"]:
                    for idx, image in enumerate(story["images"]):
                        if image:
                            st.image(image, caption=f"Ilustración {idx + 1} del cuento", use_column_width=True)
                        else:
                            st.write(f"No se pudo generar la ilustración {idx + 1}.")
                else:
                    st.write("No se pudieron generar las ilustraciones para este cuento.")

# Pie de página
st.markdown("---")
st.markdown("© 2024 Generador de Cuentos Infantiles. Desarrollado con OpenRouter, Together.xyz y Streamlit.")

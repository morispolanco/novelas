import streamlit as st
import requests
import time
import re
from io import StringIO, BytesIO
from docx import Document

st.set_page_config(
    page_title="Regenerador de Novelas",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.title("Regenerador de Novelas con OpenRouter AI")
st.write("""
    Sube tu novela en formato `.txt` o `.docx`, proporciona análisis y recomendaciones, y la aplicación regenerará tu novela capítulo por capítulo.
""")

def split_into_chapters(text):
    chapters = re.split(r'(Capítulo\s+\d+|Chapter\s+\d+)', text, flags=re.IGNORECASE)
    combined = []
    for i in range(1, len(chapters), 2):
        title = chapters[i]
        content = chapters[i+1] if i+1 < len(chapters) else ""
        combined.append(f"{title}\n{content}")
    return combined

def call_openrouter_api(prompt):
    api_url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}",
    }
    payload = {
        "model": "openai/gpt-4o-mini",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    }
    response = requests.post(api_url, headers=headers, json=payload)
    if response.status_code == 200:
        return response.json()['choices'][0]['message']['content']
    else:
        st.error(f"Error en la API: {response.status_code} - {response.text}")
        return None

def extract_text_from_docx(file):
    doc = Document(file)
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    return '\n'.join(full_text)

def create_docx(text):
    doc = Document()
    for line in text.split('\n\n'):
        doc.add_paragraph(line)
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

with st.form("regeneration_form"):
    uploaded_file = st.file_uploader("Sube tu novela (.txt o .docx)", type=["txt", "docx"])
    analysis = st.text_area("Pega tus análisis y recomendaciones aquí")
    submit_button = st.form_submit_button(label="Regenerar Novela")

if submit_button:
    if uploaded_file is not None and analysis.strip() != "":
        if uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            novel_text = extract_text_from_docx(uploaded_file)
        elif uploaded_file.type == "text/plain":
            novel_text = uploaded_file.read().decode('utf-8')
        else:
            st.error("Formato de archivo no soportado.")
            novel_text = ""

        if novel_text:
            chapters = split_into_chapters(novel_text)
            num_chapters = len(chapters)
            st.write(f"Se han detectado {num_chapters} capítulos.")

            progress_bar = st.progress(0)
            progress_text = st.empty()

            regenerated_chapters = []
            for idx, chapter in enumerate(chapters, 1):
                progress_text.text(f"Regenerando capítulo {idx} de {num_chapters}...")
                prompt = f"""
Aquí está el capítulo {idx} de mi novela:

{chapter}

Basándote en el siguiente análisis y recomendaciones, regenera este capítulo mejorándolo sin agregar comentarios o explicaciones adicionales:

{analysis}
"""
                regenerated = call_openrouter_api(prompt)
                if regenerated:
                    regenerated_chapters.append(regenerated.strip())
                else:
                    st.error("Hubo un problema al regenerar el capítulo. El proceso se detendrá.")
                    break
                progress_bar.progress(idx / num_chapters)
                time.sleep(1)

            if len(regenerated_chapters) == num_chapters:
                st.success("Novela regenerada con éxito.")
                regenerated_novel = "\n\n".join(regenerated_chapters)
                
                st.markdown("### Novela Regenerada")
                with st.expander("Ver Novela Regenerada"):
                    st.text_area("", regenerated_novel, height=400)

                st.download_button(
                    label="Descargar Novela Regenerada (.txt)",
                    data=regenerated_novel,
                    file_name="novela_regenerada.txt",
                    mime="text/plain"
                )

                docx_file = create_docx(regenerated_novel)
                st.download_button(
                    label="Descargar Novela Regenerada (.docx)",
                    data=docx_file,
                    file_name="novela_regenerada.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    else:
        st.error("Por favor, sube una novela y proporciona análisis y recomendaciones.")

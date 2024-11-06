import streamlit as st
import requests
import time
from docx import Document
from io import BytesIO
import pickle
import os

# Configuraci칩n de la p치gina
st.set_page_config(
    page_title="游닄 Generador de Libros",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.title("游닄 Generador de Libros")
st.write("Esta aplicaci칩n genera un libro basado en la idea que ingreses, dividido en cap칤tulos con t칤tulos, evitando la repetici칩n de contenido.")

# Ruta del archivo donde se guardar치 el estado
ESTADO_ARCHIVO = 'estado_generacion.pkl'

# Definir caracter칤sticas para diferentes g칠neros o tipos de libro
CARACTERISTICAS_LIBRO = {
    "Novela Juvenil": """
    **Caracter칤sticas de una buena novela juvenil:**
    1. **Extensi칩n**
       - **Longitud moderada**: Entre 40,000 y 80,000 palabras. Adaptar la extensi칩n de cada cap칤tulo para alcanzar la longitud total deseada.
    2. **Estilo**
       - **Lenguaje accesible**: Directo y sencillo, reflejando el mundo juvenil sin ser condescendiente.
       - **Narraci칩n en primera o tercera persona**: Para una conexi칩n 칤ntima con los personajes o para m칰ltiples perspectivas.
       - **Di치logos aut칠nticos**: Realistas y cre칤bles, reflejando la comunicaci칩n cotidiana de los j칩venes.
    3. **Tema**
       - **Problemas universales y espec칤ficos de la juventud**: Identidad, independencia, conflictos familiares, amistad, primer amor, presi칩n social, descubrimiento personal, salud mental, acoso, racismo, discriminaci칩n, etc.
       - **Desarrollo emocional**: Enfocado en el crecimiento emocional de los personajes, mostrando c칩mo enfrentan y superan sus miedos y limitaciones.
    4. **Protagonistas atractivos y cercanos**
       - J칩venes de edades cercanas a la audiencia, con caracter칤sticas atractivas pero imperfectas, enfrentando dilemas morales y evolucionando a lo largo de la historia.
    5. **Subg칠neros variados**
       - Romance, ciencia ficci칩n, fantas칤a, aventuras, misterio, horror, etc., manteniendo el enfoque en temas relevantes para la adolescencia.
    6. **Narrativa 치gil**
       - Ritmo r치pido para captar la atenci칩n, con cap칤tulos cortos y giros frecuentes en la trama.
    7. **Mensaje positivo o inspirador**
       - Transmitir mensajes de superaci칩n, esperanza, autenticidad, tolerancia y empat칤a, ayudando a los lectores a enfrentar sus propios desaf칤os personales.
    """,
    "Novela de Ciencia Ficci칩n": """
    **Caracter칤sticas de una buena novela de ciencia ficci칩n:**
    1. **Mundo Elaborado**
       - **Construcci칩n detallada del universo**: Reglas f칤sicas, tecnol칩gicas y sociales bien definidas.
    2. **Innovaci칩n Tecnol칩gica**
       - **Tecnolog칤as avanzadas o futuristas** que impulsan la trama y generan conflictos.
    3. **Temas Filos칩ficos y 칄ticos**
       - **Exploraci칩n de cuestiones profundas** como la inteligencia artificial, la 칠tica cient칤fica, la identidad, etc.
    4. **Personajes Complejos**
       - **Protagonistas y antagonistas** con motivaciones claras y evoluci칩n a lo largo de la historia.
    5. **Trama Intrigante**
       - **Conflictos y giros inesperados** que mantienen el inter칠s del lector.
    6. **Integraci칩n de Ciencia y Ficci칩n**
       - **Base cient칤fica s칩lida** combinada con elementos ficticios para crear una narrativa cre칤ble.
    7. **Visi칩n del Futuro**
       - **Predicciones o reflexiones** sobre el futuro de la humanidad y la tecnolog칤a.
    """,
    "Ensayo": """
    **Caracter칤sticas de un buen ensayo:**
    1. **Tesis Clara**
       - **Idea principal** bien definida que gu칤a todo el ensayo.
    2. **Estructura Coherente**
       - **Introducci칩n, desarrollo y conclusi칩n** organizados de manera l칩gica.
    3. **Argumentaci칩n S칩lida**
       - **Evidencias y ejemplos** que respaldan la tesis.
    4. **Estilo Propio**
       - **Voz 칰nica** que refleja el pensamiento y la personalidad del autor.
    5. **Profundidad y Reflexi칩n**
       - **An치lisis profundo** de los temas tratados, mostrando un entendimiento claro.
    6. **Claridad y Concisi칩n**
       - **Lenguaje preciso y directo** que facilita la comprensi칩n del lector.
    7. **Originalidad**
       - **Perspectivas innovadoras** o enfoques 칰nicos sobre el tema.
    """,
    "Autoayuda": """
    **Caracter칤sticas de un buen libro de autoayuda:**
    1. **Objetivo Claro**
       - **Enfoque espec칤fico** en mejorar aspectos concretos de la vida del lector.
    2. **Estructura Pr치ctica**
       - **Incluye ejercicios, ejemplos y consejos** aplicables a situaciones reales.
    3. **Lenguaje Motivador**
       - **Tonos positivos y alentadores** que inspiran al lector a tomar acci칩n.
    4. **Testimonios y Casos de Estudio**
       - **Historias reales** que ilustran los puntos clave y demuestran eficacia.
    5. **Conclusiones Accionables**
       - **Pasos concretos** que el lector puede seguir para implementar los consejos.
    6. **Accesibilidad**
       - **Contenido f치cil de entender** y aplicar para cualquier persona, independientemente de su trasfondo.
    7. **Inspiraci칩n y Empoderamiento**
       - **Fomenta la confianza** y la capacidad del lector para superar desaf칤os personales.
    """,
    "Psicolog칤a": """
    **Caracter칤sticas de un buen libro de psicolog칤a:**
    1. **Fundamento Cient칤fico**
       - **Basado en teor칤as y estudios** psicol칩gicos reconocidos.
    2. **An치lisis Profundo**
       - **Exploraci칩n detallada** de conceptos y fen칩menos psicol칩gicos.
    3. **Casos de Estudio**
       - **Ejemplos pr치cticos** que ilustran teor칤as y aplicaciones.
    4. **Claridad y Precisi칩n**
       - **Lenguaje t칠cnico pero accesible**, adecuado para lectores no especializados.
    5. **Aplicaciones Pr치cticas**
       - **C칩mo aplicar los conceptos** en la vida diaria para mejorar el bienestar psicol칩gico.
    6. **Estructura Coherente**
       - **Organizaci칩n l칩gica** que facilita la comprensi칩n y retenci칩n del contenido.
    7. **Reflexi칩n y Autoconocimiento**
       - **Fomenta la introspecci칩n** y el entendimiento personal del lector.
    """,
    "Relaciones": """
    **Caracter칤sticas de un buen libro sobre relaciones:**
    1. **Din치mica Interpersonal**
       - **Exploraci칩n de diferentes tipos de relaciones** (amistades, parejas, familiares).
    2. **Consejos Pr치cticos**
       - **Estrategias para mejorar la comunicaci칩n** y resolver conflictos.
    3. **Desarrollo Emocional**
       - **Fomento de la inteligencia emocional** y la empat칤a.
    4. **Historias y Ejemplos**
       - **Narrativas que ejemplifican conceptos clave** y facilitan la comprensi칩n.
    5. **Reflexi칩n Personal**
       - **Ejercicios para que el lector eval칰e** y mejore sus propias relaciones.
    6. **Diversidad y Inclusi칩n**
       - **Consideraci칩n de diferentes perspectivas** y experiencias en las relaciones.
    7. **Construcci칩n de Confianza**
       - **M칠todos para establecer y mantener la confianza** en las relaciones interpersonales.
    """,
    "Negocios": """
    **Caracter칤sticas de un buen libro de negocios:**
    1. **Estrategias Empresariales**
       - **M칠todos para iniciar, gestionar y escalar** negocios de manera efectiva.
    2. **An치lisis de Mercado**
       - **T칠cnicas para investigar y comprender** el mercado objetivo y la competencia.
    3. **Gesti칩n Financiera**
       - **Fundamentos de finanzas empresariales**, incluyendo presupuestos, inversiones y flujo de caja.
    4. **Liderazgo y Gesti칩n de Equipos**
       - **Desarrollo de habilidades de liderazgo** efectivo y gesti칩n de equipos de alto rendimiento.
    5. **Innovaci칩n y Competitividad**
       - **Fomento de la creatividad** y adaptaci칩n al cambio para mantener la competitividad.
    6. **Planificaci칩n Estrat칠gica**
       - **Definici칩n de objetivos claros** y desarrollo de planes para alcanzarlos.
    7. **Casos de 칄xito y Fracaso**
       - **Ejemplos reales** que ilustran mejores pr치cticas y lecciones aprendidas.
    """
    # Puedes agregar m치s tipos seg칰n sea necesario
}

def guardar_estado():
    """Guarda el estado de la sesi칩n en un archivo local."""
    with open(ESTADO_ARCHIVO, 'wb') as f:
        pickle.dump({
            'capitulos': st.session_state.capitulos,
            'resumenes': st.session_state.resumenes,
            'titulo_obra': st.session_state.titulo_obra,
            'proceso_generado': st.session_state.proceso_generado,
            'prompt': st.session_state.prompt,
            'tipo_libro': st.session_state.tipo_libro
        }, f)

def cargar_estado():
    """Carga el estado de la sesi칩n desde un archivo local."""
    if os.path.exists(ESTADO_ARCHIVO):
        with open(ESTADO_ARCHIVO, 'rb') as f:
            estado = pickle.load(f)
            st.session_state.capitulos = estado.get('capitulos', [])
            st.session_state.resumenes = estado.get('resumenes', [])
            st.session_state.titulo_obra = estado.get('titulo_obra', "Libro")
            st.session_state.proceso_generado = estado.get('proceso_generado', False)
            st.session_state.prompt = estado.get('prompt', "")
            st.session_state.tipo_libro = estado.get('tipo_libro', list(CARACTERISTICAS_LIBRO.keys())[0])
        return True
    return False

def limpiar_estado():
    """Limpia el estado de la sesi칩n y elimina el archivo de estado si existe."""
    st.session_state.capitulos = []
    st.session_state.resumenes = []
    st.session_state.titulo_obra = "Libro"
    st.session_state.proceso_generado = False
    st.session_state.prompt = ""
    st.session_state.tipo_libro = list(CARACTERISTICAS_LIBRO.keys())[0]
    if os.path.exists(ESTADO_ARCHIVO):
        os.remove(ESTADO_ARCHIVO)

# Inicializar estado de la sesi칩n
if 'capitulos' not in st.session_state:
    st.session_state.capitulos = []
if 'resumenes' not in st.session_state:
    st.session_state.resumenes = []
if 'titulo_obra' not in st.session_state:
    st.session_state.titulo_obra = "Libro"
if 'proceso_generado' not in st.session_state:
    st.session_state.proceso_generado = False
if 'prompt' not in st.session_state:
    st.session_state.prompt = ""
if 'tipo_libro' not in st.session_state:
    st.session_state.tipo_libro = list(CARACTERISTICAS_LIBRO.keys())[0]  # Valor por defecto

# Intentar cargar el estado guardado al iniciar la aplicaci칩n
estado_cargado = cargar_estado()

# Funci칩n para generar un cap칤tulo de cualquier tipo de libro
def generar_capitulo(prompt, capitulo_num, resumen_previas, tipo_libro):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
    }
    instrucciones = (
        "Aseg칰rate de que el contenido generado cumpla con las caracter칤sticas del tipo de libro seleccionado. "
        "Desarrolla los conceptos clave y explora sus aplicaciones. "
        "Utiliza un lenguaje adecuado al g칠nero, desarrolla ideas relevantes y cercanas a la audiencia, "
        "aborda temas pertinentes para el tipo de libro y mantiene una narrativa coherente con el g칠nero. "
        "Evita repetir informaci칩n ya mencionada en cap칤tulos anteriores. "
        "Cada cap칤tulo debe comenzar con un t칤tulo apropiado."
    )
    if resumen_previas:
        resumen_texto = " Hasta ahora, el libro ha cubierto los siguientes puntos: " + resumen_previas
    else:
        resumen_texto = ""
    
    caracteristicas = CARACTERISTICAS_LIBRO.get(tipo_libro, "")
    
    mensaje = (
        f"{caracteristicas}\n\n"
        f"Escribe el cap칤tulo {capitulo_num} de un libro de tipo '{tipo_libro}' sobre el siguiente tema: {prompt}. "
        f"El cap칤tulo debe comenzar con un t칤tulo apropiado y tener aproximadamente 2000 palabras. "
        f"No debe contener subdivisiones ni subcap칤tulos.{resumen_texto} {instrucciones}"
    )
    data = {
        "model": "openai/gpt-4o-mini",
        "messages": [
            {
                "role": "user",
                "content": mensaje
            }
        ]
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        respuesta = response.json()
        if 'choices' in respuesta and len(respuesta['choices']) > 0:
            contenido_completo = respuesta['choices'][0]['message']['content']
            lineas = contenido_completo.strip().split('\n', 1)
            if len(lineas) == 2:
                titulo_capitulo = lineas[0].strip().replace("T칤tulo:", "").replace("Titulo:", "").strip()
                contenido = lineas[1].strip()
                return titulo_capitulo, contenido
            else:
                st.warning(f"No se pudo extraer el t칤tulo del Cap칤tulo {capitulo_num}.")
                return f"Cap칤tulo {capitulo_num}", contenido_completo
        else:
            st.error(f"Respuesta inesperada de la API al generar el cap칤tulo {capitulo_num}.")
            return None, None
    except requests.exceptions.RequestException as e:
        st.error(f"Error al generar el cap칤tulo {capitulo_num}: {e}")
        return None, None

# Funci칩n para resumir un cap칤tulo utilizando la API de OpenRouter
def resumir_capitulo(capitulo, tipo_libro):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
    }
    caracteristicas = CARACTERISTICAS_LIBRO.get(tipo_libro, "")
    prompt_resumen = (
        f"{caracteristicas}\n\n"
        "Proporciona un resumen conciso y relevante del siguiente cap칤tulo del libro. "
        "El resumen debe resaltar los puntos clave de la trama, los desarrollos de los conceptos y los eventos principales, "
        "evitando detalles redundantes.\n\n"
        f"Cap칤tulo:\n{capitulo}\n\nResumen:"
    )
    data = {
        "model": "openai/gpt-4o-mini",
        "messages": [
            {
                "role": "user",
                "content": prompt_resumen
            }
        ]
    }
    try:
        response = requests.post(url, headers=headers, json=data)
        response.raise_for_status()
        respuesta = response.json()
        if 'choices' in respuesta and len(respuesta['choices']) > 0:
            resumen = respuesta['choices'][0]['message']['content']
            resumen = ' '.join(resumen.split())
            return resumen
        else:
            st.error("Respuesta inesperada de la API al resumir el cap칤tulo.")
            return None
    except requests.exceptions.RequestException as e:
        st.error(f"Error al resumir el cap칤tulo: {e}")
        return None

# Funci칩n para crear el documento Word con t칤tulos
def crear_documento(capitulo_list, titulo, tipo_libro):
    doc = Document()
    doc.add_heading(titulo, 0)
    doc.add_paragraph(f"Tipo de Libro: {tipo_libro}")
    doc.add_paragraph()
    for idx, (titulo_capitulo, capitulo) in enumerate(capitulo_list, 1):
        doc.add_heading(f"Cap칤tulo {idx}: {titulo_capitulo}", level=1)
        doc.add_paragraph(capitulo)
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Interfaz de usuario para seleccionar opciones
st.sidebar.title("Opciones")

# Determinar las opciones disponibles en la barra lateral
opciones_disponibles = []
if estado_cargado and len(st.session_state.capitulos) < 24:
    opciones_disponibles = ["Continuar Generando", "Iniciar Nueva Generaci칩n"]
else:
    opciones_disponibles = ["Iniciar Nueva Generaci칩n"]

# Radio buttons sin necesidad de bot칩n de env칤o
opcion = st.sidebar.radio("쯈u칠 deseas hacer?", opciones_disponibles)

mostrar_formulario = False
if opcion == "Iniciar Nueva Generaci칩n":
    limpiar_estado()
    mostrar_formulario = True
elif opcion == "Continuar Generando":
    if len(st.session_state.capitulos) >= 24:
        st.sidebar.info("Has alcanzado el l칤mite m치ximo de 24 cap칤tulos.")
    else:
        mostrar_formulario = True

if mostrar_formulario:
    with st.form(key='form_libro'):
        if opcion == "Iniciar Nueva Generaci칩n":
            st.session_state.tipo_libro = st.selectbox(
                "Selecciona el tipo de libro que deseas generar:",
                options=list(CARACTERISTICAS_LIBRO.keys())
            )
            st.session_state.prompt = st.text_area(
                "Ingresa la idea o tema para el libro:",
                height=200,
                value=""
            )
        else:
            st.selectbox(
                "Tipo de libro:",
                options=list(CARACTERISTICAS_LIBRO.keys()),
                index=list(CARACTERISTICAS_LIBRO.keys()).index(st.session_state.tipo_libro),
                disabled=True
            )
            st.text_area(
                "Idea o tema para el libro:",
                height=200,
                value=st.session_state.prompt,
                disabled=True
            )
        
        cap_generadas = len(st.session_state.capitulos)
        cap_restantes = 24 - cap_generadas
        num_capitulos = st.slider(
            "N칰mero de cap칤tulos a generar:",
            min_value=1,
            max_value=cap_restantes,
            value=min(3, cap_restantes)
        )
        submit_button = st.form_submit_button(label='Generar Libro')
    
    if submit_button:
        if opcion == "Iniciar Nueva Generaci칩n":
            if not st.session_state.prompt.strip():
                st.error("Por favor, ingresa una idea o tema v치lida para el libro.")
                st.stop()
        else:
            pass
        
        st.success("Iniciando la generaci칩n del libro...")
        st.session_state.proceso_generado = True
        progreso = st.progress(0)
        
        inicio = len(st.session_state.capitulos) + 1
        fin = inicio + num_capitulos - 1
        if fin > 24:
            fin = 24
        cap_generadas_en_ejecucion = 0
        
        for i in range(inicio, fin + 1):
            st.write(f"Generando **Cap칤tulo {i}**...")
            if st.session_state.resumenes:
                resumen_previas = ' '.join(st.session_state.resumenes)
            else:
                resumen_previas = ''
            titulo_capitulo, capitulo = generar_capitulo(
                st.session_state.prompt, 
                i, 
                resumen_previas,
                st.session_state.tipo_libro
            )
            if capitulo:
                st.session_state.capitulos.append((titulo_capitulo, capitulo))
                resumen = resumir_capitulo(capitulo, st.session_state.tipo_libro)
                if resumen:
                    st.session_state.resumenes.append(resumen)
                else:
                    st.warning(f"No se pudo generar un resumen para el Cap칤tulo {i}.")
                guardar_estado()
                cap_generadas_en_ejecucion += 1
            else:
                st.error("La generaci칩n del libro se ha detenido debido a un error.")
                break
            progreso.progress(cap_generadas_en_ejecucion / num_capitulos)
            time.sleep(2)
        
        progreso.empty()
        
        if cap_generadas_en_ejecucion == num_capitulos:
            st.success(f"Se han generado {cap_generadas_en_ejecucion} cap칤tulos exitosamente.")
            st.session_state.titulo_obra = st.text_input("T칤tulo del libro:", value=st.session_state.titulo_obra)
            if st.session_state.titulo_obra:
                documento = crear_documento(st.session_state.capitulos, st.session_state.titulo_obra, st.session_state.tipo_libro)
                st.download_button(
                    label="Descargar Libro en Word",
                    data=documento,
                    file_name="libro.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
        else:
            st.info(f"Generaci칩n interrumpida. Has generado {cap_generadas_en_ejecucion} de {num_capitulos} cap칤tulos.")

# Mostrar el libro generado
if st.session_state.capitulos and st.session_state.proceso_generado:
    st.markdown("---")
    st.header("游닀 Libro Generado")
    for idx, (titulo_capitulo, capitulo) in enumerate(st.session_state.capitulos, 1):
        st.subheader(f"Cap칤tulo {idx}: {titulo_capitulo}")
        st.write(capitulo)

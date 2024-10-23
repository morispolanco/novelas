import streamlit as st
import requests
from io import BytesIO
import re
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
import json

# =====================
# Configuración Inicial
# =====================

OPENROUTER_API_KEY = st.secrets["OPENROUTER_API_KEY"]
API_URL = "https://openrouter.ai/api/v1/chat/completions"
MODEL = "openai/gpt-4o-mini"  # Asegúrate de que este es el nombre correcto del modelo

# =====================
# Funciones Auxiliares
# =====================

def generar_contenido(prompt, max_tokens=2000, temperature=0.7, repetition_penalty=1.2, frequency_penalty=0.5):
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {OPENROUTER_API_KEY}"
    }
    data = {
        "model": MODEL,
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ],
        "max_tokens": max_tokens,
        "temperature": temperature,
        "repetition_penalty": repetition_penalty,
        "frequency_penalty": frequency_penalty
    }
    try:
        response = requests.post(API_URL, headers=headers, json=data, timeout=120)  # Aumentar timeout si es necesario
        response.raise_for_status()
        return response.json()["choices"][0]["message"]["content"]
    except requests.exceptions.Timeout:
        st.error("La solicitud a la API de OpenRouter ha excedido el tiempo de espera.")
        return ""
    except requests.exceptions.ConnectionError:
        st.error("Error de conexión al intentar comunicarse con la API de OpenRouter.")
        return ""
    except requests.exceptions.HTTPError as e:
        st.error(f"Error HTTP de la API de OpenRouter: {e.response.status_code} - {e.response.text}")
        return ""
    except (KeyError, IndexError) as e:
        st.error("Respuesta inesperada de la API de OpenRouter.")
        st.error(str(e))
        return ""
    except Exception as e:
        st.error(f"Ocurrió un error inesperado: {e}")
        return ""

def exportar_a_docx(contenido_novela, titulo="Novela Mejorada"):
    doc = Document()
    try:
        # Configuración de la página
        sections = doc.sections
        for section in sections:
            section.page_width = Inches(6)
            section.page_height = Inches(9)
            section.top_margin = Inches(0.75)
            section.bottom_margin = Inches(0.75)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)

        # Configuración de la fuente predeterminada a Alegreya, tamaño 12
        style_normal = doc.styles['Normal']
        font_normal = style_normal.font
        font_normal.name = 'Alegreya'
        font_normal.size = Pt(12)
        style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), 'Alegreya')

        # Configurar fuentes para los estilos de encabezado
        for heading in ['Heading 1', 'Heading 2', 'Heading 3']:
            style_heading = doc.styles[heading]
            font_heading = style_heading.font
            font_heading.name = 'Alegreya'
            font_heading.size = Pt(14)
            style_heading.element.rPr.rFonts.set(qn('w:eastAsia'), 'Alegreya')

        # Título de la novela
        doc.add_heading(titulo, level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph("\n")  # Espacio después del título

        # Procesar el contenido
        for linea in contenido_novela.split('\n'):
            linea = linea.strip()
            if not linea:
                continue

            if re.match(r"^Escena\s+\d+", linea, re.IGNORECASE):
                p = doc.add_heading(linea, level=2)
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)
            else:
                p = doc.add_paragraph(linea, style='Normal')
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.line_spacing = 1.15
                p.paragraph_format.space_after = Pt(6)

        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    except Exception as e:
        st.error(f"Error al formatear el documento DOCX: {e}")
        return None

def leer_docx(file):
    try:
        doc = Document(file)
        full_text = []
        for para in doc.paragraphs:
            full_text.append(para.text)
        return '\n'.join(full_text)
    except Exception as e:
        st.error(f"Error al leer el archivo DOCX: {e}")
        return ""

def dividir_en_escenas(texto_novela):
    """
    Divide el texto de la novela en escenas basándose en delimitadores específicos.
    Por ejemplo, se espera que cada escena comience con "Escena X" donde X es un número.
    """
    try:
        # Regex para detectar el inicio de una escena
        escenas = re.split(r'(?i)(Escena\s+\d+)', texto_novela)
        # Combina los encabezados de escena con su contenido
        escenas_combinadas = []
        for i in range(1, len(escenas), 2):
            encabezado = escenas[i].strip()
            contenido = escenas[i+1].strip() if i+1 < len(escenas) else ""
            escenas_combinadas.append(f"{encabezado}\n{contenido}")
        return escenas_combinadas
    except Exception as e:
        st.error(f"Error al dividir la novela en escenas: {e}")
        return []

# =====================
# Cacheo de Funciones
# =====================

@st.cache_data(show_spinner=False, ttl=3600)
def generar_contenido_cache(prompt, max_tokens=2000, temperature=0.7, repetition_penalty=1.2, frequency_penalty=0.5):
    return generar_contenido(prompt, max_tokens, temperature, repetition_penalty, frequency_penalty)

# =====================
# Validación de Entrada
# =====================

def validar_correcciones(correcciones):
    if not correcciones:
        return False, "Las correcciones y mejoras no pueden estar vacías."
    if len(correcciones) < 10:
        return False, "Las correcciones y mejoras son demasiado cortas. Por favor, proporciona más detalles."
    if len(correcciones) > 2000:
        return False, "Las correcciones y mejoras son demasiado largas. Por favor, proporciona una descripción más concisa."
    return True, ""

# =====================
# Inicialización de Variables de Sesión
# =====================

if 'novela_mejorada' not in st.session_state:
    st.session_state.novela_mejorada = False  # Bandera para evitar reprocesamiento

# =====================
# Interfaz de Usuario
# =====================

st.set_page_config(page_title="Mejorar tu Novela", layout="wide")

# Título principal
st.title("Mejorar tu Novela Escena por Escena")

# ============================================
# Sección: Mejorar una Novela Subida por el Usuario
# ============================================

st.header("Mejorar tu Novela")

# Paso 1: Subir el archivo DOCX
st.subheader("Paso 1: Sube tu Novela en Formato DOCX")
archivo_subido = st.file_uploader("Selecciona tu archivo DOCX:", type=["docx"])

# Paso 2: Especificar las Correcciones y Mejoras
st.subheader("Paso 2: Especifica las Correcciones y Mejoras")
correcciones = st.text_area("Describe los defectos que has identificado en tu novela y las mejoras que deseas aplicar a cada escena:", height=200)

# Botón para iniciar la mejora
if st.button("Aplicar Correcciones y Mejoras", key="aplicar_mejoras"):
    if not archivo_subido:
        st.error("Por favor, sube un archivo DOCX para poder proceder.")
    else:
        correcciones_validas, mensaje_error = validar_correcciones(correcciones)
        if not correcciones_validas:
            st.error(mensaje_error)
        else:
            with st.spinner("Procesando tu novela y aplicando las correcciones y mejoras..."):
                # Leer el contenido del archivo DOCX
                contenido_novela = leer_docx(archivo_subido)
                if not contenido_novela.strip():
                    st.error("El archivo DOCX está vacío o no se pudo extraer el contenido.")
                else:
                    # Dividir la novela en escenas
                    escenas = dividir_en_escenas(contenido_novela)
                    if not escenas:
                        st.error("No se encontraron escenas delimitadas en la novela. Asegúrate de que cada escena comience con 'Escena X' donde X es un número.")
                    else:
                        escenas_mejoradas = []
                        total_escenas = len(escenas)
                        tareas_completadas = 0

                        progress_bar = st.progress(0)
                        status_text = st.empty()

                        for escena in escenas:
                            # Extraer el encabezado de la escena
                            encabezado = escena.split('\n')[0]
                            contenido = '\n'.join(escena.split('\n')[1:]).strip()

                            # Crear el prompt para mejorar la escena
                            prompt_mejora = (
                                f"Corrige los defectos señalados y aplica las mejoras sugeridas en la siguiente escena de una novela. "
                                f"Defectos y mejoras a aplicar:\n{correcciones}\n\n"
                                f"Texto de la escena:\n{contenido}\n\n"
                                f"Por favor, reescribe la escena incorporando las correcciones y mejoras mencionadas. "
                                f"Mantén la coherencia, el estilo y el tono original de la obra."
                            )

                            # Generar el contenido mejorado de la escena
                            contenido_mejorado = generar_contenido_cache(
                                prompt_mejora,
                                max_tokens=2000,  # Ajusta según sea necesario
                                temperature=0.7,
                                repetition_penalty=1.2,
                                frequency_penalty=0.5
                            )

                            if contenido_mejorado:
                                # Añadir el encabezado y el contenido mejorado
                                escena_mejorada = f"{encabezado}\n{contenido_mejorado}"
                                escenas_mejoradas.append(escena_mejorada)
                            else:
                                # Si hubo un error, conservar la escena original y notificar al usuario
                                escenas_mejoradas.append(escena)
                                st.warning(f"No se pudo mejorar la escena: {encabezado}")

                            tareas_completadas += 1
                            status_text.text(f"Mejorada Escena {tareas_completadas} de {total_escenas}")
                            progress_bar.progress(tareas_completadas / total_escenas)

                        # Unir todas las escenas mejoradas en una sola novela
                        novela_mejorada = '\n\n'.join(escenas_mejoradas)

                        if novela_mejorada:
                            st.session_state.contenido_mejorado = novela_mejorada
                            st.session_state.novela_mejorada = True
                            st.success("Tu novela ha sido mejorada exitosamente.")
                        else:
                            st.error("Ocurrió un error al intentar mejorar tu novela.")

# Mostrar el contenido mejorado y permitir la descarga
if 'contenido_mejorado' in st.session_state and st.session_state.novela_mejorada:
    st.subheader("Paso 3: Revisa y Descarga tu Novela Mejorada")
    contenido_editable_mejorado = st.text_area(
        "Revisa y edita el contenido mejorado si es necesario:", 
        st.session_state.contenido_mejorado, 
        height=800
    )

    # Botón para descargar el archivo mejorado
    if st.button("Descargar Novela Mejorada", key="descargar_mejorada"):
        buffer_docx_mejorada = exportar_a_docx(contenido_editable_mejorado, titulo="Novela Mejorada")
        if buffer_docx_mejorada:
            st.download_button(
                label="Descargar Novela Mejorada en DOCX",
                data=buffer_docx_mejorada,
                file_name="novela_mejorada.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
            st.success("Descarga iniciada correctamente.")
        else:
            st.error("Ocurrió un error al generar el archivo DOCX.")

    # Botón para resetear la sección de mejora
    if st.button("Mejorar Otra Novela", key="reset_mejorar_novela"):
        for key in ['contenido_mejorado', 'novela_mejorada']:
            if key in st.session_state:
                del st.session_state[key]
        st.experimental_rerun()

# =====================
# Finalización
# =====================

st.markdown("""
---
*Nota: Asegúrate de que el archivo DOCX que subas esté bien estructurado con delimitadores claros para cada escena (por ejemplo, "Escena 1", "Escena 2", etc.). Esto facilitará el proceso de división y reescritura escena por escena.*
""")

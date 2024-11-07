import streamlit as st
import requests
from docx import Document
from io import BytesIO
import backoff
import re
from difflib import SequenceMatcher
import pandas as pd

# Definir la cantidad máxima de capítulos
MAX_CAPITULOS = 24

# Configuración de la página
st.set_page_config(
    page_title="📝 Adventure Tales Generator for Kids (9-12 years)",
    layout="wide",
    initial_sidebar_state="expanded",
)

st.title("📝 Adventure Tales Generator for Kids (9-12 years)")
st.write("""
Esta aplicación genera hasta 24 capítulos de cuentos de aventuras para niños de 9 a 12 años en inglés.
El usuario proporciona una Tabla de Contenidos con resúmenes, y la aplicación generará cada cuento entre 500 y 700 palabras.
Cada capítulo comienza con la palabra "CHAPTER".
""")

# Inicializar estado de la sesión
if 'capitulos' not in st.session_state:
    st.session_state.capitulos = []
if 'titulo_obra' not in st.session_state:
    st.session_state.titulo_obra = "Adventure Tales"
if 'proceso_generado' not in st.session_state:
    st.session_state.proceso_generado = False
if 'temas_utilizados' not in st.session_state:
    st.session_state.temas_utilizados = []

# Función para medir la similitud entre dos textos
def similar(a, b):
    return SequenceMatcher(None, a.lower(), b.lower()).ratio()

# Prompt personalizado refinado
PROMPT_BASE = """
Write an adventure story intended for boys and girls aged between 9 and 12 years. The story should be exciting and age-appropriate, including elements such as challenges, brave characters, and imaginative settings. Ensure the content is entertaining, but also safe and suitable for children. Include an interesting conflict and a resolution that leaves a positive message.

# Requirements and Suggestions
- The story should be between 500 and 700 words, using accessible and understandable language for readers in this age group.
- Introduce lovable characters that readers can empathize with, such as curious children, magical animals, or fantastical beings.
- There should be at least one obstacle or challenge that the characters must overcome, with a positive message about teamwork, bravery, or creativity at the end.
- Use visual descriptions to create vibrant scenes, but avoid using overly complex terms or situations.

# Suggested Structure
1. **Introduction**: Present the protagonist and the initial setting where a calm situation is experienced before the adventure begins.
2. **Conflict**: An event that changes the protagonist's routine and leads them to an unexpected mission.
3. **Development**: Moments of action where the protagonist faces challenges and obstacles. There can be a companion who helps the protagonist.
4. **Resolution**: The conclusion of the adventure with a creative solution and a happy ending that offers reflection or a positive message.

# Tone and Style
- **Tone**: Adventurous, motivating, fun.
- **Narrative Style**: Third person or first person.

# Output Format
The output should include:
1. **CHAPTER {n}: {Title}**
2. **Summary:** A brief summary of the chapter.
3. **Story Content:** The full story, between 500-700 words.

Each time a speaking character changes, use a line break for clarity.

# Chapter Title Format
Begin the story with "CHAPTER {n}: {Title}", where {n} is the chapter number and {Title} is the story's title.

# Unique Theme Instruction
Each chapter must have a unique theme that has not been used in previous chapters. Refer to the list of used themes below and choose a new, distinct theme for this story. Avoid using similar phrases or titles such as "The Quest for...", "The Mystery of...", or "The Adventure of...".
"""

# Función para extraer el título usando expresiones regulares
def extraer_titulo(respuesta, capitulo_num):
    # Buscamos "CHAPTER {n}: Título"
    patron = rf'CHAPTER\s*{capitulo_num}:\s*(.*)'
    match = re.search(patron, respuesta, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return "Title Not Found"

# Función para extraer el resumen del cuento
def extraer_resumen(respuesta):
    # Buscar la línea que comienza con "Summary:"
    patron = r'Summary:\s*(.*)'
    match = re.search(patron, respuesta, re.IGNORECASE)
    if match:
        return match.group(1).strip()
    return "Summary Not Found"

# Función para extraer el tema principal del cuento
def extraer_tema(respuesta):
    # Buscar la línea que comienza con "Summary:" y extraer la última oración como tema
    patron = r'Summary:\s*(.*)'
    match = re.search(patron, respuesta, re.IGNORECASE)
    if match:
        resumen = match.group(1).strip()
        # Dividir en oraciones y tomar la última
        oraciones = re.split(r'[.!?]', resumen)
        if len(oraciones) >= 1:
            return oraciones[-1].strip()
    return "Unique Theme Not Found"

# Función con reintentos para generar un capítulo
@backoff.on_exception(backoff.expo, requests.exceptions.RequestException, max_tries=3)
def generar_capitulo(capitulo_num, titulo, resumen):
    url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
    }

    # Construir el prompt incluyendo la lista de temas utilizados
    temas_utilizados = st.session_state.temas_utilizados
    if temas_utilizados:
        temas_formateados = "; ".join(temas_utilizados)
        prompt = PROMPT_BASE.replace("Refer to the list of used themes below and choose a new, distinct theme for this story.",
                                    f"Refer to the list of used themes below and choose a new, distinct theme for this story.\n\nUsed Themes: {temas_formateados}")
    else:
        prompt = PROMPT_BASE.replace("Refer to the list of used themes below and choose a new, distinct theme for this story.",
                                    "Refer to the list of used themes below and choose a new, distinct theme for this story.\n\nUsed Themes: None")

    mensaje = (
        f"{prompt}\n\n"
        f"CHAPTER {capitulo_num}: {titulo}\n"
        f"Summary: {resumen}\n"
        f"Story Content:"
    )

    data = {
        "model": "openai/gpt-4",  # Cambiar al modelo correcto si es necesario
        "messages": [
            {
                "role": "user",
                "content": mensaje
            }
        ]
    }

    response = requests.post(url, headers=headers, json=data)
    response.raise_for_status()
    respuesta = response.json()
    if 'choices' in respuesta and len(respuesta['choices']) > 0:
        contenido_completo = respuesta['choices'][0]['message']['content']
        titulo_generado = extraer_titulo(contenido_completo, capitulo_num)
        resumen_generado = extraer_resumen(contenido_completo)
        tema_generado = extraer_tema(contenido_completo)
        # Extraer el contenido sin el título y el resumen
        contenido = contenido_completo.replace(f"CHAPTER {capitulo_num}: {titulo_generado}", "", 1)
        contenido = contenido.replace(f"Summary: {resumen_generado}", "", 1).strip()
        return titulo_generado, resumen_generado, contenido, tema_generado
    else:
        st.error(f"Unexpected API response when generating Chapter {capitulo_num}.")
        return None, None, None, None

# Función para crear el documento Word con tabla de contenidos y capítulos
def crear_documento(capitulos_list, titulo):
    doc = Document()
    doc.add_heading(titulo, 0)
    
    # Crear Tabla de Contenidos
    doc.add_heading("Table of Contents", level=1)
    for idx, (titulo_capitulo, resumen_capitulo, _, _) in enumerate(capitulos_list, 1):
        toc_entry = f"CHAPTER {idx}: {titulo_capitulo}"
        toc_summary = f"Summary: {resumen_capitulo}"
        doc.add_paragraph(toc_entry, style='List Number')
        doc.add_paragraph(toc_summary, style='List Bullet')
    
    doc.add_page_break()
    
    # Agregar cada capítulo
    for idx, (titulo_capitulo, _, capitulo, _) in enumerate(capitulos_list, 1):
        doc.add_heading(f"CHAPTER {idx}: {titulo_capitulo}", level=1)
        doc.add_paragraph(capitulo)
        doc.add_page_break()
    
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# Función para crear un DataFrame de la Tabla de Contenidos con Resúmenes
def crear_tabla_contenidos(capitulos_list):
    data = {
        "Chapter": [f"CHAPTER {idx}" for idx in range(1, len(capitulos_list)+1)],
        "Title": [capitulo[0] for capitulo in capitulos_list],
        "Summary": [capitulo[1] for capitulo in capitulos_list]
    }
    df = pd.DataFrame(data)
    return df

# Interfaz de usuario para subir un archivo CSV o ingresar manualmente
st.sidebar.title("Options")

opciones_disponibles = ["Upload CSV File", "Manual Input"]
opcion = st.sidebar.radio("How would you like to provide the Table of Contents?", opciones_disponibles)

mostrar_formulario = False
if opcion == "Upload CSV File":
    archivo = st.sidebar.file_uploader("Upload a CSV file with columns: Chapter, Title, Summary", type=["csv"])
    if archivo:
        try:
            df = pd.read_csv(archivo)
            if not all(col in df.columns for col in ["Chapter", "Title", "Summary"]):
                st.error("El archivo CSV debe contener las columnas: Chapter, Title, Summary")
            else:
                st.session_state.capitulos_input = df
                mostrar_formulario = True
        except Exception as e:
            st.error(f"Error reading the CSV file: {e}")
else:
    mostrar_formulario = True

if opcion == "Manual Input" or (opcion == "Upload CSV File" and 'capitulos_input' in st.session_state):
    with st.form(key='form_cuento_infantil'):
        if opcion == "Manual Input":
            num_capitulos_input = st.number_input(
                "Número de capítulos a generar:",
                min_value=1,
                max_value=MAX_CAPITULOS,
                value=3
            )
            capitulos_manual = []
            for i in range(1, num_capitulos_input + 1):
                st.markdown(f"### Chapter {i}")
                titulo = st.text_input(f"Title for Chapter {i}", key=f"title_{i}")
                resumen = st.text_area(f"Summary for Chapter {i}", key=f"summary_{i}", height=100)
                capitulos_manual.append({"Chapter": f"CHAPTER {i}", "Title": titulo, "Summary": resumen})
        else:
            df = st.session_state.capitulos_input
            capitulos_manual = df.to_dict('records')
        
        submit_button = st.form_submit_button(label='Generate Adventure Tales')

    if submit_button:
        if opcion == "Manual Input":
            capitulos = capitulos_manual
            if any(not cap['Title'] or not cap['Summary'] for cap in capitulos):
                st.error("Por favor, asegúrate de que todos los capítulos tengan un título y un resumen.")
            else:
                st.session_state.capitulos_input = pd.DataFrame(capitulos)
        else:
            capitulos = capitulos_manual
        
        st.success("Starting the generation of adventure tales...")
        st.session_state.proceso_generado = True
        progreso = st.progress(0)
        
        capitulos_generados_ejecucion = 0
        total_capitulos = len(capitulos)
        
        for idx, capitulo in enumerate(capitulos, 1):
            titulo_input = capitulo["Title"]
            resumen_input = capitulo["Summary"]
            st.write(f"Generating **CHAPTER {idx}: {titulo_input}**...")
            titulo_generado, resumen_generado, contenido, tema_generado = generar_capitulo(idx, titulo_input, resumen_input)
            if contenido and tema_generado:
                # Verificar similitud con temas existentes
                repetido = False
                for tema in st.session_state.temas_utilizados:
                    if similar(tema_generado, tema) > 0.7:  # Umbral de similitud
                        repetido = True
                        break
                if not repetido:
                    st.session_state.capitulos.append((titulo_generado, resumen_generado, contenido, tema_generado))
                    st.session_state.temas_utilizados.append(tema_generado)
                    capitulos_generados_ejecucion += 1
                else:
                    st.warning(f"The theme '{tema_generado}' is too similar to an existing theme. Attempting to generate a new chapter...")
                    # Intentar re-generar el capítulo
                    titulo_generado_new, resumen_generado_new, contenido_new, tema_generado_new = generar_capitulo(idx, titulo_input, resumen_input)
                    if contenido_new and tema_generado_new and similar(tema_generado_new, tema_generado) < 0.7:
                        st.session_state.capitulos.append((titulo_generado_new, resumen_generado_new, contenido_new, tema_generado_new))
                        st.session_state.temas_utilizados.append(tema_generado_new)
                        capitulos_generados_ejecucion += 1
                    else:
                        st.error("Could not generate a chapter with a unique theme after multiple attempts.")
                        break
            else:
                st.error("Generation of tales has been stopped due to an error.")
                break
            progreso.progress(capitulos_generados_ejecucion / total_capitulos)
            # time.sleep(1)  # Removed to improve speed
        
        progreso.empty()
        
        if capitulos_generados_ejecucion == total_capitulos:
            st.success(f"Successfully generated {capitulos_generados_ejecucion} chapters.")
            st.session_state.titulo_obra = st.text_input("Title of the adventure tales:", value=st.session_state.titulo_obra)
            if st.session_state.titulo_obra:
                documento = crear_documento(st.session_state.capitulos, st.session_state.titulo_obra)
                st.download_button(
                    label="Download Tales in Word",
                    data=documento,
                    file_name="adventure_tales.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
                # Optional: Add PDF download
                # documento_pdf = crear_pdf(st.session_state.capitulos, st.session_state.titulo_obra)
                # st.download_button(
                #     label="Download Tales in PDF",
                #     data=documento_pdf,
                #     file_name="adventure_tales.pdf",
                #     mime="application/pdf"
                # )
        else:
            st.info(f"Generation interrupted. You have generated {capitulos_generados_ejecucion} out of {total_capitulos} chapters.")

# Mostrar los capítulos generados
if st.session_state.capitulos and st.session_state.proceso_generado:
    st.markdown("---")
    st.header("📖 Generated Adventure Tales")
    
    # Mostrar Tabla de Contenidos
    st.subheader("Table of Contents")
    for idx, (titulo_capitulo, resumen_capitulo, _, _) in enumerate(st.session_state.capitulos, 1):
        st.markdown(f"**CHAPTER {idx}: {titulo_capitulo}**")
        st.markdown(f"*Summary:* {resumen_capitulo}")
        st.markdown("---")
    
    # Mostrar cada capítulo
    for idx, (titulo_capitulo, _, capitulo, _) in enumerate(st.session_state.capitulos, 1):
        st.markdown(f"### CHAPTER {idx}: {titulo_capitulo}")
        st.write(capitulo)
        st.markdown("---")

import streamlit as st
import json
from PIL import Image
from io import BytesIO
import requests
import base64
from docx import Document
from docx.shared import Inches

# Configuración de la página
st.set_page_config(
    page_title="Biblioteca de Cuentos Clásicos con Ilustraciones",
    layout="centered",
    initial_sidebar_state="expanded",
)

# Verificar que las claves API estén configuradas
if not st.secrets.get('TOGETHER_API_KEY'):
    st.sidebar.error("Por favor, configura la clave API de Together.xyz en los secretos de Streamlit.")
    st.stop()

# Cargar historias desde el archivo JSON
@st.cache_data
def load_stories():
    with open('stories.json', 'r', encoding='utf-8') as file:
        return json.load(file)

stories_data = load_stories()

# Título de la aplicación
st.title("📚 Biblioteca de Cuentos Clásicos con Ilustraciones")

# Descripción
st.markdown("""
Explora y disfruta de una colección de **Fábulas de Esopo**, **Cuentos de La Fontaine**, **Cuentos de los Hermanos Grimm** y **Cuentos de Hans Christian Andersen**. 
Selecciona las historias que deseas leer y obtener ilustraciones personalizadas.
""")

# Barra lateral para seleccionar las historias
st.sidebar.header("Selecciona las Historias")

# Obtener las categorías (autores)
categories = list(stories_data.keys())

# Crear un diccionario para almacenar las selecciones
selected_stories = {}

for category in categories:
    titles = [story['título'] for story in stories_data[category]]
    selected = st.sidebar.multiselect(
        label=f"{category}",
        options=titles,
        key=category
    )
    selected_stories[category] = selected

# Botón para generar ilustraciones
generate_images = st.sidebar.button("Generar Ilustraciones")

# Variable para almacenar las historias seleccionadas
stories_to_display = []

# Iterar sobre las selecciones y recopilar las historias
for category, titles in selected_stories.items():
    for title in titles:
        # Encontrar la historia en el JSON
        story = next((s for s in stories_data[category] if s['título'] == title), None)
        if story:
            stories_to_display.append({
                "categoría": category,
                "título": story['título'],
                "texto": story['texto']
            })

# Función para generar una ilustración usando Together.xyz
def generate_image(prompt_description):
    together_api_key = st.secrets.get('TOGETHER_API_KEY')
    if not together_api_key:
        st.error("La clave API de Together.xyz no está configurada en los secretos de Streamlit.")
        return None

    api_url = "https://api.together.xyz/v1/images/generations"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {together_api_key}"
    }

    data = {
        "model": "black-forest-labs/FLUX.1.1-pro",
        "prompt": prompt_description,
        "width": 512,
        "height": 512,
        "steps": 50,  # Ajusta según la calidad deseada
        "n": 1,
        "response_format": "b64_json"
    }

    try:
        response = requests.post(api_url, headers=headers, data=json.dumps(data))
        response.raise_for_status()
        result = response.json()
        b64_image = result['data'][0]['b64_json']
        image_bytes = base64.b64decode(b64_image)
        image = Image.open(BytesIO(image_bytes))
        return image
    except requests.exceptions.HTTPError as http_err:
        st.error(f"Error HTTP al generar la imagen: {http_err}")
        if 'response' in locals():
            st.write(response.text)  # Mostrar la respuesta completa para depuración
    except Exception as err:
        st.error(f"Error al generar la imagen: {err}")
    return None

# Si hay historias seleccionadas y el usuario ha presionado el botón para generar ilustraciones
if stories_to_display and generate_images:
    st.markdown("### Historias Seleccionadas con Ilustraciones")
    doc = Document()
    doc.add_heading("Biblioteca de Cuentos Clásicos con Ilustraciones", 0)

    for idx, story in enumerate(stories_to_display, start=1):
        st.markdown(f"#### {story['título']} ({story['categoría']})")
        st.write(story['texto'])

        # Generar descripciones para las ilustraciones basadas en la historia
        image_prompt1 = f"Una ilustración colorida y atractiva para un cuento infantil sobre '{story['título']}'. Representa un momento clave de la historia."
        image_prompt2 = f"Una segunda ilustración que represente otro momento significativo en la historia de '{story['título']}'."

        # Generar las imágenes
        image1 = generate_image(image_prompt1)
        image2 = generate_image(image_prompt2)

        # Mostrar las imágenes en la aplicación
        if image1:
            st.image(image1, caption="Ilustración 1", use_column_width=True)
        else:
            st.write("No se pudo generar la ilustración 1.")

        if image2:
            st.image(image2, caption="Ilustración 2", use_column_width=True)
        else:
            st.write("No se pudo generar la ilustración 2.")

        st.markdown("---")

        # Añadir al documento de Word
        doc.add_heading(f"{story['título']}", level=1)
        doc.add_paragraph(story['texto'])

        # Añadir imágenes al documento
        for img in [image1, image2]:
            if img:
                img_byte_arr = BytesIO()
                img.save(img_byte_arr, format='PNG')
                img_byte_arr = img_byte_arr.getvalue()
                doc.add_picture(BytesIO(img_byte_arr), width=Inches(4))
                doc.add_paragraph("")  # Espacio adicional después de la imagen
            else:
                doc.add_paragraph("No se pudo generar la ilustración.")

    # Guardar el documento en un flujo de BytesIO
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)

    # Proporcionar el botón de descarga
    st.success("¡Historias e ilustraciones generadas con éxito!")
    st.download_button(
        label="📄 Descargar Historias como Documento de Word",
        data=doc_io,
        file_name="Biblioteca_Cuentos_Clásicos.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Mostrar las historias seleccionadas sin ilustraciones
if stories_to_display:
    st.markdown("### Historias Seleccionadas")
    for story in stories_to_display:
        st.markdown(f"#### {story['título']} ({story['categoría']})")
        st.write(story['texto'])
        st.markdown("---")

# Pie de página
st.markdown("---")
st.markdown("© 2024 Biblioteca de Cuentos Clásicos. Todos los derechos reservados.")

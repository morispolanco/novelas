import streamlit as st
import requests
from docx import Document
import os
import time

st.set_page_config(page_title="Generador de Capítulos", layout="wide")

st.title("Generador de Capítulos a partir de una Tabla de Contenidos")

toc = st.text_area("Introduce la tabla de contenidos (uno por línea):", height=200)

if st.button("Generar Capítulos"):
    if toc.strip() == "":
        st.error("Por favor, introduce una tabla de contenidos.")
    else:
        capitulos = toc.strip().split('\n')
        documento = Document()
        progreso = st.progress(0)
        total = len(capitulos)
        api_key = os.getenv("OPENROUTER_API_KEY")
        
        if not api_key:
            st.error("Clave de API no encontrada. Por favor, configura OPENROUTER_API_KEY.")
        else:
            for idx, capitulo in enumerate(capitulos, 1):
                st.write(f"### Generando: {capitulo}")
                
                payload = {
                    "model": "openai/gpt-4o-mini",
                    "messages": [
                        {
                            "role": "user",
                            "content": f"Escribe el contenido para el capítulo titulado '{capitulo}'."
                        }
                    ]
                }
                
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}"
                }
                
                response = requests.post(
                    "https://openrouter.ai/api/v1/chat/completions",
                    headers=headers,
                    json=payload
                )
                
                if response.status_code == 200:
                    data = response.json()
                    texto = data.get("choices")[0].get("message").get("content").strip()
                    documento.add_heading(capitulo, level=1)
                    documento.add_paragraph(texto)
                else:
                    documento.add_heading(capitulo, level=1)
                    documento.add_paragraph("Error al generar el contenido.")
                
                progreso.progress(idx / total)
                time.sleep(1)
            
            documento.save("Capitulos.docx")
            
            with open("Capitulos.docx", "rb") as f:
                st.download_button(
                    label="Descargar Documento Word",
                    data=f,
                    file_name="Capitulos.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            
            st.success("¡Capítulos generados exitosamente!")

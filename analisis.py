import streamlit as st
import requests
import json
from docx import Document
from io import BytesIO
import re
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from datetime import datetime

# Configuración de la página
st.set_page_config(
    page_title="Analizador de Novelas de Suspenso Político",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Título de la aplicación
st.title("Analizador de Novelas de Suspenso Político")
st.write("""
Esta aplicación analiza una novela en el género de thriller político.
Sube tu novela en formato `.docx` o `.txt` y recibe un informe detallado de errores y mejoras, además de una calificación global.
""")

# Barra lateral para opciones de usuario
st.sidebar.header("Configuración de Análisis")
file_upload = st.sidebar.file_uploader("Sube tu novela (.docx o .txt):", type=["docx", "txt"])

# Estado de la aplicación
if 'etapa' not in st.session_state:
    st.session_state.etapa = "inicio"
if 'novela' not in st.session_state:
    st.session_state.novela = ""
if 'informe' not in st.session_state:
    st.session_state.informe = ""

def call_openrouter_api(prompt, max_tokens=3000, temperature=0.7, top_p=0.9, top_k=50, repetition_penalty=1.2):
    """
    Llama a la API de OpenRouter para obtener una respuesta basada en el prompt proporcionado.
    
    Parámetros:
        prompt (str): El texto de entrada para enviar a la API.
        max_tokens (int): Número máximo de tokens en la respuesta.
        temperature (float): Parámetro de temperatura para la generación.
        top_p (float): Parámetro top_p para la generación.
        top_k (int): Parámetro top_k para la generación.
        repetition_penalty (float): Penalización por repetición.
    
    Retorna:
        str or None: La respuesta de la API o None si ocurre un error.
    """
    api_url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "openai/gpt-4o-mini",  # Asegúrate de que este modelo sea compatible
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": max_tokens,
        "top_p": top_p,
        "top_k": top_k,
        "repetition_penalty": repetition_penalty,
        "stop": ["[\"<|eot_id|>\"]"],
        "stream": False
    }
    
    session = requests.Session()
    retries = Retry(total=5, backoff_factor=1, status_forcelist=[502, 503, 504])
    session.mount('https://', HTTPAdapter(max_retries=retries))
    
    try:
        response = session.post(api_url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
        response_json = response.json()
        if 'choices' in response_json and len(response_json['choices']) > 0:
            return response_json['choices'][0]['message']['content']
        else:
            st.error("La respuesta de la API no contiene 'choices'.")
            st.write("Respuesta completa de la API:", response_json)
            return None
    except requests.exceptions.RequestException as e:
        st.error(f"Error en la llamada a la API: {e}")
        return None

def leer_archivo(file):
    """
    Lee el contenido del archivo subido, soportando .docx y .txt.
    
    Parámetros:
        file: El archivo subido por el usuario.
    
    Retorna:
        str or None: El texto del archivo o None si ocurre un error.
    """
    try:
        if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            # Leer archivo .docx
            doc = Document(file)
            texto = "\n".join([para.text for para in doc.paragraphs])
        elif file.type == "text/plain":
            # Leer archivo .txt
            texto = file.read().decode("utf-8")
        else:
            st.error("Formato de archivo no soportado.")
            return None
        return texto
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")
        return None

def dividir_en_escenas(texto):
    """
    Divide el texto de la novela en escenas basadas en dos saltos de línea.
    
    Parámetros:
        texto (str): El texto completo de la novela.
    
    Retorna:
        list: Lista de escenas divididas.
    """
    # Suponiendo que las escenas están separadas por dos saltos de línea
    escenas = texto.split('\n\n')
    escenas = [escena.strip() for escena in escenas if escena.strip()]
    return escenas

def analizar_novela(texto):
    """
    Analiza la novela enviándola a la API y obteniendo un análisis detallado, incluyendo recomendaciones específicas por escena.
    
    Parámetros:
        texto (str): El texto completo de la novela.
    
    Retorna:
        str or None: El análisis obtenido de la API o None si falla.
    """
    escenas = dividir_en_escenas(texto)
    mejoras_por_escena = []

    for idx, escena in enumerate(escenas, start=1):
        prompt = f"""
        Realiza una valoración literaria completa de la siguiente escena de una novela de suspenso político. 
        Analiza aspectos como coherencia, consistencia, desarrollo de personajes, profundidad psicológica, ritmo, estilo, descripciones, diálogos, vocabulario y cualquier otro aspecto relevante. 
        Señala los puntos que deben mejorar o cambiar para cada aspecto mencionado. 
        Además, asigna una puntuación de 1 a 10 para esta escena basada en su calidad literaria.
        Presenta la respuesta en formato JSON con las siguientes claves: "escena", "valoracion", "mejoras", "puntuacion".
        
        Escena {idx}:
        {escena}
        
        ### Informe de Análisis:
        """
        analisis = call_openrouter_api(prompt)
        if analisis:
            try:
                analisis_json = json.loads(analisis)
                mejoras_por_escena.append({
                    "escena": analisis_json.get("escena", f"Escena {idx}"),
                    "valoracion": analisis_json.get("valoracion", "No disponible."),
                    "mejoras": analisis_json.get("mejoras", "No se identificaron puntos de mejora."),
                    "puntuacion": analisis_json.get("puntuacion", "N/A")
                })
            except json.JSONDecodeError:
                st.error(f"Error al decodificar la respuesta JSON para la Escena {idx}.")
                continue
        else:
            st.error(f"No se pudo obtener análisis para la Escena {idx}.")

    # Una vez analizadas todas las escenas, generar el informe global
    prompt_global = f"""
    Basándote en las valoraciones y mejoras específicas de cada escena, realiza una valoración literaria global de la novela de suspenso político. 
    Identifica los principales errores gramaticales, de coherencia, desarrollo de personajes, ritmo, estilo, etc. 
    Proporciona recomendaciones claras y específicas para mejorar la calidad general de la novela. 
    Además, asigna una puntuación global de 1 a 10 puntos basada en la calidad literaria de la novela.
    Presenta la respuesta en formato JSON con las siguientes claves: "calificacion", "errores", "recomendaciones", "puntuacion".
    
    ### Mejoras por Escena:
    {json.dumps(mejoras_por_escena, ensure_ascii=False)}
    
    ### Informe de Análisis Global:
    """
    analisis_global = call_openrouter_api(prompt_global)
    return analisis_global

def generar_informe(analisis):
    """
    Genera un informe formateado a partir del análisis JSON obtenido de la API, incluyendo mejoras específicas por escena.
    
    Parámetros:
        analisis (str): La respuesta JSON de la API.
    
    Retorna:
        str: El informe formateado o None si falla.
    """
    if not analisis:
        st.error("No hay análisis para generar el informe.")
        return None
    try:
        # Intentar cargar la respuesta como JSON
        analisis_json = json.loads(analisis)
        # Validar que las claves existen
        required_keys = {"calificacion", "errores", "recomendaciones", "puntuacion"}
        if not required_keys.issubset(analisis_json.keys()):
            st.error("La respuesta JSON de la API está incompleta.")
            return None
        calificacion = analisis_json.get('calificacion', 'No disponible')
        errores = analisis_json.get('errores', 'No se identificaron errores.')
        recomendaciones = analisis_json.get('recomendaciones', 'No se proporcionaron recomendaciones.')
        puntuacion = analisis_json.get('puntuacion', 'N/A')
        
        informe = f"# Informe de Análisis de la Novela\n\n"
        informe += f"**Calificación General:** {calificacion} / 10\n\n"
        informe += f"**Puntuación Literaria:** {puntuacion} / 10\n\n"
        informe += f"**Errores Identificados:**\n{errores}\n\n"
        informe += f"**Recomendaciones para Mejoras:**\n{recomendaciones}\n\n"
        informe += f"## Valoración Específica por Escena\n\n"
        
        for mejora in analisis_json.get('mejoras_por_escena', []):
            informe += f"### {mejora['escena']}\n\n"
            informe += f"**Valoración:** {mejora['valoracion']}\n\n"
            informe += f"**Puntos de Mejora:** {mejora['mejoras']}\n\n"
            informe += f"**Puntuación de la Escena:** {mejora['puntuacion']} / 10\n\n"
        
        return informe
    except json.JSONDecodeError:
        # Si la respuesta no es JSON, mostrar el contenido completo
        informe = f"# Informe de Análisis de la Novela\n\nCalificación General: No disponible\n\nErrores y Recomendaciones:\n{analisis}\n\n"
        return informe

def exportar_informe_word(informe):
    """
    Exporta el informe de análisis a un documento Word, incluyendo mejoras específicas por escena.
    
    Parámetros:
        informe (str): El informe formateado.
    
    Retorna:
        BytesIO: El buffer del documento Word.
    """
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    # Establecer el estilo normal con una fuente común
    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    # Agregar el título
    titulo_paragraph = document.add_heading("Informe de Análisis de la Novela", level=1)
    titulo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_paragraph()  # Añadir un párrafo en blanco

    # Agregar el contenido del informe
    for line in informe.split('\n'):
        if line.startswith("# "):
            # Títulos
            document.add_heading(line.replace("# ", ""), level=2)
        elif line.startswith("## "):
            # Subtítulos
            document.add_heading(line.replace("## ", ""), level=3)
        elif line.startswith("### "):
            # Subtítulos de escena
            document.add_heading(line.replace("### ", ""), level=4)
        elif line.startswith("**"):
            # Negritas
            match = re.match(r'\*\*(.*?)\*\*:\s*(.*)', line)
            if match:
                term, desc = match.groups()
                p = document.add_paragraph()
                p.add_run(f"{term}: ").bold = True
                p.add_run(desc)
        else:
            # Párrafos normales
            if line.strip():
                document.add_paragraph(line)

    # Guardar el documento en memoria
    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

def mostrar_inicio():
    """
    Muestra la interfaz de inicio donde el usuario puede cargar su novela.
    """
    st.header("Carga y Análisis de la Novela")
    if file_upload:
        texto = leer_archivo(file_upload)
        if texto:
            st.session_state.novela = texto
            st.session_state.etapa = "analisis"

def mostrar_analisis():
    """
    Muestra la interfaz de análisis donde el usuario puede iniciar el análisis de la novela.
    """
    st.header("Análisis en Proceso")
    novela = st.session_state.novela

    # Mostrar un extracto de la novela para confirmar
    st.subheader("Extracto de la Novela:")
    extracto = novela[:1000] + "..." if len(novela) > 1000 else novela
    st.text_area("Contenido de la Novela:", extracto, height=200, disabled=True)

    # Botón para iniciar el análisis
    enviar = st.button("Enviar")

    if enviar:
        # Crear una barra de progreso
        progress_bar = st.progress(0)
        status_text = st.empty()

        try:
            # Paso 1: Analizar la novela
            status_text.text("Iniciando el análisis de la novela...")
            progress_bar.progress(10)
            analisis = analizar_novela(novela)
            if not analisis:
                st.error("No se pudo generar el análisis. Por favor, intenta nuevamente.")
                progress_bar.progress(0)
                return

            progress_bar.progress(50)

            # Paso 2: Generar el informe
            status_text.text("Generando el informe...")
            informe = generar_informe(analisis)
            if not informe:
                st.error("No se pudo generar el informe de análisis.")
                progress_bar.progress(0)
                return

            progress_bar.progress(80)

            # Paso 3: Finalizar y mostrar completado
            st.session_state.informe = informe
            st.session_state.etapa = "completado"
            progress_bar.progress(100)
            status_text.text("Análisis completado exitosamente.")

        except Exception as e:
            st.error(f"Ocurrió un error durante el análisis: {e}")
            progress_bar.progress(0)
            status_text.text("Error durante el análisis.")

def mostrar_completado():
    """
    Muestra el informe de análisis generado y proporciona opciones para descargarlo.
    """
    st.header("Informe de Análisis Generado")
    informe = st.session_state.informe

    # Mostrar el informe en la interfaz
    st.subheader("Informe Detallado:")
    st.markdown(informe)  # Usar markdown para una mejor visualización

    # Exportar a Word
    doc_buffer = exportar_informe_word(informe)
    st.download_button(
        label="Descargar Informe en Word",
        data=doc_buffer,
        file_name=f"informe_analisis_novela_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Interfaz de usuario principal
if st.session_state.etapa == "inicio":
    mostrar_inicio()

elif st.session_state.etapa == "analisis":
    mostrar_analisis()

elif st.session_state.etapa == "completado":
    mostrar_completado()

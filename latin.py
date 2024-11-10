import streamlit as st
import requests
import json
from docx import Document
from docx.shared import Inches
from io import BytesIO
import random
from PIL import Image
import base64

# Configuración de la página
st.set_page_config(
    page_title="Generador de Cuentos Infantiles en Latín Simplificado con Ilustraciones",
    layout="centered",
    initial_sidebar_state="auto",
)

# Título de la aplicación
st.title("📚 Generador de Cuentos Infantiles en Latín Simplificado con Ilustraciones")

# Descripción
st.markdown("""
Genera cuentos personalizados en **latín simplificado** para niños, adaptados a diferentes grupos de edad. 
Especifica el número de cuentos (hasta 15), y deja que la aplicación cree historias atractivas con temas variados, longitudes apropiadas y dos hermosas ilustraciones por cuento.
""")

# Lista predefinida de 150 temas
themes = [
    "aventura en el bosque",
    "un viaje mágico",
    "amistad y trabajo en equipo",
    "superar miedos",
    "un día en la playa",
    "exploración espacial",
    "misterio submarino",
    "viaje en el tiempo",
    "una búsqueda heroica",
    "aprender a compartir",
    "el castillo encantado",
    "una historia de detectives",
    "compañeros robots",
    "tierra de fantasía",
    "reino animal",
    "el dragón amigable",
    "la princesa valiente",
    "el robot perdido",
    "la isla secreta",
    "un monstruo en el armario",
    "la fábrica de chocolate",
    "el tren mágico",
    "los juguetes que cobran vida",
    "una aventura en el desierto",
    "el bosque de los sueños",
    "el misterioso castillo",
    "el héroe inesperado",
    "la escuela de magia",
    "el cohete a la luna",
    "el puente arcoíris",
    "el viaje al fondo del mar",
    "la casa de los árboles",
    "el robot y el niño",
    "la feria encantada",
    "el tesoro escondido",
    "la carrera de globos",
    "el bosque de las hadas",
    "el caballo volador",
    "la carrera espacial",
    "el jardín secreto",
    "la tormenta de estrellas",
    "el mago y el aprendiz",
    "la ciudad perdida",
    "el misterio del lago",
    "la aventura en la montaña",
    "el libro de los deseos",
    "la isla de los dinosaurios",
    "el viaje al centro de la tierra",
    "el circo de los sueños",
    "el unicornio mágico",
    "la amistad con un extraterrestre",
    "aventuras en el mundo de los sueños",
    "el tesoro del viejo faro",
    "la carrera contra el tiempo",
    "el misterio de la cueva escondida",
    "viaje al mundo de los insectos",
    "la ciudad flotante",
    "el guardián del bosque",
    "el portal secreto",
    "la fiesta de cumpleaños sorpresa",
    "el reino de los colores",
    "la historia del pequeño gigante",
    "el misterio de las estrellas perdidas",
    "la magia de la lluvia",
    "el viaje al país de los juguetes",
    "la aventura en el planeta desconocido",
    "el dragón de hielo",
    "la travesía del pequeño barco",
    "el jardín encantado",
    "la amistad entre una bruja y un hada",
    "el misterio de la luna llena",
    "la búsqueda del arcoíris perdido",
    "el mundo invertido",
    "la isla de los sueños olvidados",
    "el lobo sabio",
    "la torre mágica",
    "el príncipe y el sapo",
    "la sirena valiente",
    "el dragón de fuego",
    "la estrella fugaz",
    "el bosque de los duendes",
    "la isla de los tesoros",
    "el viaje al país de las maravillas",
    "la cabaña misteriosa",
    "el castillo de hielo",
    "la heroína escondida",
    "el mago oscuro",
    "la nave perdida",
    "el reino de las nubes",
    "la princesa del desierto",
    "el puente encantado",
    "la aventura en el volcán",
    "el jardín de las mariposas",
    "la luna de los deseos",
    "el valle de los sueños",
    "la mariposa dorada",
    "el unicornio de la montaña",
    "la historia del árbol milenario",
    "el guardián del arcoíris",
    "la cueva de los misterios",
    "el castillo de los espejos",
    "la princesa guerrera",
    "el río de la magia",
    "la luna llena y el lobo",
    "el reino submarino",
    "la aventura en la selva tropical",
    "el mago de los vientos",
    "la dama del lago",
    "el viaje al país de los gigantes",
    "la ciudad de las luces",
    "el secreto del bosque",
    "la sirena del arrecife",
    "el héroe del desierto",
    "la princesa de las estrellas",
    "el dragón y el niño",
    "la luna de plata",
    "el viaje al corazón del bosque",
    "la estrella del norte",
    "el reino de las sombras",
    "la aventura en el cielo",
    "el mago del tiempo",
    "la princesa del mar",
    "el castillo de los dragones",
    "la historia del lobo solitario",
    "el viaje a la tierra de los sueños",
    "la princesa y el dragón",
    "el guardián de la montaña",
    "la magia del bosque",
    "el viaje al reino de cristal",
    "la sirena y el príncipe",
    "el mago de la selva",
    "la aventura en el mundo invertido",
    "el castillo de los sueños",
    "la luna y las estrellas",
    "el viaje al país de los cuentos",
    "la princesa y el unicornio",
    "el guardián de la cueva",
    "la historia del hada olvidada",
    "el dragón de la tormenta",
    "la aventura en el jardín secreto",
    "el mago de la oscuridad",
    "la sirena y el lobo",
    "el viaje al reino de las hadas",
    "la princesa de los vientos",
    "el guardián de los sueños",
    "la magia de la luna",
    "el viaje al corazón del océano",
    "la historia del lobo y la luna",
    "el castillo de las estrellas",
    "la princesa y el mago",
    "el dragón de la selva",
    "la aventura en el reino de las nubes",
    "el mago de los mares",
    "la sirena de las profundidades",
    "el viaje al país de los sueños",
    "la princesa del bosque encantado",
    "el guardián del río",
    "la historia del hada y el unicornio",
    "el dragón y la princesa",
    "la aventura en el valle mágico",
    "el mago del bosque"
    ]

    # Lista predefinida de nombres de personajes únicos
character_names = [
        "Luna", "Mateo", "Sofía", "Emilio", "Valentina", "Lucas", "Isabella", "Alejandro",
        "Camila", "Gabriel", "Emma", "Benjamín", "Victoria", "Daniel", "Lucía", "Diego",
        "Martina", "Samuel", "Natalia", "Sebastián", "Valeria", "Emiliano", "Catalina",
        "Amelia", "Jorge", "Renata", "Andrés", "Sara", "Antonio", "Claudia",
        "Pablo", "Mía", "Ricardo", "Alicia", "Javier", "Paula", "Santiago", "Gabriela",
        "Hugo", "María", "Fernando", "Julia", "Adrián", "Lorena", "Tomás", "Andrea",
        "Óscar", "Fernanda"
    ]

    # Función para generar un único nombre de personaje
    def get_unique_name(used_names):
        available_names = list(set(character_names) - set(used_names))
        if not available_names:
            # Si se agotan los nombres, generar uno aleatorio
            return f"Personaje_{random.randint(1000, 9999)}"
        return random.choice(available_names)

    # Barra lateral para la entrada del usuario
    st.sidebar.header("Parámetros de Entrada")

    # Entrada para el número de cuentos
    num_stories = st.sidebar.number_input(
        "Número de Cuentos",
        min_value=1,
        max_value=15,
        value=1,
        step=1
    )

    # Grupos de edad para elegir
    age_groups = ["3-5 años", "6-8 años", "9-12 años"]
    selected_age_group = st.sidebar.selectbox(
        "Selecciona el Grupo de Edad para los Cuentos",
        age_groups
    )

    # Botón para generar cuentos
    if st.sidebar.button("Generar Cuentos"):
        with st.spinner("Generando cuentos e ilustraciones..."):
            # Función para generar y corregir un solo cuento usando OpenRouter
            def generate_corrected_story(theme, age_group, character_name):
                api_url = "https://openrouter.ai/api/v1/chat/completions"
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
                }

                # Definir las 20 directrices
                guidelines = """
    1. **Desarrollo de Personajes:**
       - Profundiza en los conflictos internos de los personajes. Muestra sus miedos, dudas y pequeñas victorias para hacerlos más humanos y relacionables. A medida que cada personaje enfrenta un desafío, incluye momentos donde reflexionen sobre sus inseguridades o limitaciones para enriquecer su crecimiento.

    2. **Consistencia en el Tono y Edad:**
       - Las historias están dirigidas a lectores de 12 años, pero algunos temas y emociones pueden ser ligeramente avanzados para esta edad. Simplifica ciertas decisiones morales y emocionales, o haz que los personajes actúen y piensen de manera apropiada para su edad para mantener la autenticidad.

    3. **Detalles Sensoriales:**
       - Utiliza descripciones sensoriales para hacer que los entornos y experiencias sean más inmersivos. Incorpora olores, sonidos y texturas para crear una atmósfera más rica y ayudar a los lectores a sentirse más conectados con el mundo de la historia.

    4. **Morales Sutiles:**
       - Presenta las lecciones de cada historia de manera más implícita en lugar de explícita. Permite que los personajes y los lectores descubran la moral a través de acciones y consecuencias, haciendo que el mensaje sea más efectivo y menos moralizante.

    5. **Profundización de Relaciones entre Personajes:**
       - Los lazos entre personajes añaden una capa emocional a las historias. Muestra interacciones que revelen más sobre sus relaciones (conflictos, apoyo, amistades) para ayudar a los lectores a sentirse más conectados emocionalmente con la historia.

    6. **Conflictos y Desafíos Graduales:**
       - Asegúrate de que los personajes enfrenten obstáculos en una progresión natural. Esto añade tensión y hace que la resolución sea más satisfactoria, ya que el personaje supera diversos desafíos para lograr su objetivo.

    7. **Variedad de Conflictos:**
       - Varía los tipos de conflictos (internos, externos, naturales, sobrenaturales, emocionales) para dar dinamismo a la colección y evitar un patrón repetitivo. Esto permite la exploración de diferentes tipos de coraje, empatía y creatividad.

    8. **Símbolos y Elementos Recurrentes:**
       - Incluye símbolos o elementos recurrentes (como objetos, frases o lugares) que tengan un significado especial en cada historia. Esto unifica la colección temáticamente, ayudando a los lectores a encontrar conexiones y significados a lo largo de las historias.

    9. **Contexto del Mundo y Antecedentes:**
       - En historias con un mundo mágico o fantástico, ofrece un breve contexto del entorno o las "reglas mágicas" para construir un universo consistente. Esto mejora el atractivo y la inmersión de la historia.

    10. **Mostrar, No Contar:**
        - En lugar de decir directamente al lector lo que un personaje siente o aprende, muéstralo a través de acciones, gestos y decisiones. Esta técnica narrativa permite que los lectores interpreten las emociones del personaje y extraigan la moral por sí mismos.

    11. **Fortalecer Temas de Fondo:**
        - Asegúrate de que cada historia tenga un tema central claro (como valentía, amistad o empatía) y explóralo en profundidad. Los temas consistentes a lo largo de la colección le dan una identidad cohesiva.

    12. **Desarrollar Personajes Memorables:**
        - Los personajes principales deben tener cualidades distintivas (como valentía, ingenio o humildad) que se reflejen en sus acciones y decisiones. Añade pequeños detalles físicos o emocionales para hacerlos más humanos y fáciles de recordar para los lectores.

    13. **Agregar Detalles Sensoriales:**
        - Utiliza descripciones sensoriales (olor, tacto, sonido) para sumergir a los lectores en el entorno de cada historia. Estos detalles enriquecen la narrativa y ayudan a los lectores a sentir que están explorando estos mundos con los personajes.

    14. **Incorporar Conflictos o Desafíos Claros:**
        - Cada historia debe tener un conflicto o desafío central para que el protagonista lo supere. Esto crea tensión y mantiene a los lectores enganchados. Además, resolver problemas o enfrentar dilemas hace que las historias sean más dinámicas.

    15. **Mostrar Crecimiento y Aprendizaje:**
        - Asegúrate de que los personajes crezcan o aprendan algo significativo en cada historia. Las lecciones o cambios internos deben sentirse naturales y reflejarse en la narrativa sin parecer forzados, haciendo que el aprendizaje sea algo que los lectores puedan sentir y comprender.

    16. **Crear una Conexión con el Lector:**
        - Utiliza situaciones o emociones que los niños puedan relacionar, como la curiosidad, el miedo a lo desconocido o la búsqueda de aceptación. Esto les ayuda a sentirse parte de la historia y a comprender mejor el mensaje.

    17. **Añadir Humor y Toques Ligeros:**
        - Aunque las historias deben contener una lección, añade humor en diálogos o interacciones entre personajes. Esto no solo hace que la lectura sea más entretenida, sino que también equilibra la historia, especialmente si el mensaje es profundo.

    18. **Usar Ritmo y Variedad en la Narrativa:**
        - Juega con el ritmo, alternando momentos de calma con picos de emoción o tensión. Además, varía el lenguaje y el estilo de narración en cada historia para mantener la frescura, incluso con un tono general consistente.

    19. **Crear Introducciones y Conclusiones Memorables:**
        - Las líneas de apertura deben captar la atención del lector, y las conclusiones deben dejar una sensación de satisfacción o reflexión. Esto hace que la historia sea más impactante y memorable.

    20. **Transmitir Valores Positivos sin Ser Moralista:**
        - En lugar de hacer la moral explícita, permite que los valores positivos emerjan naturalmente de las decisiones y acciones de los personajes. Esto evita que las historias se sientan moralizantes y da al mensaje una resonancia poderosa y sutil.
    """

                # Añadir las directrices al prompt
                prompt = f"""
    {guidelines}

    Crea una historia que cumpla con las siguientes características:

    - Tema: {theme}
    - Grupo de edad: {age_group}
    - Nombre del personaje principal: {character_name}
    - Texto continuo sin subdivisiones ni subtítulos.
    - Escribe la historia en **latín simplificado**, utilizando una gramática y vocabulario adecuados para niños que están aprendiendo el idioma.
    """

                data = {
                    "model": "gpt-4",  # Asegúrate de que este modelo esté disponible en OpenRouter
                    "messages": [
                        {
                            "role": "user",
                            "content": prompt
                        }
                    ],
                    "temperature": 0.7,
                    "max_tokens": 1500  # Aumentar tokens para asegurar suficiente contenido en latín
                }

                try:
                    response = requests.post(api_url, headers=headers, data=json.dumps(data))
                    response.raise_for_status()
                    result = response.json()

                    # Verificar estructura de la respuesta
                    if 'choices' in result and len(result['choices']) > 0:
                        story = result['choices'][0]['message']['content'].strip()
                    else:
                        st.error("Respuesta inesperada de la API de OpenRouter.")
                        return "Lo siento, ocurrió un error al generar el cuento."

                    return story
                except requests.exceptions.HTTPError as http_err:
                    st.error(f"Ocurrió un error HTTP: {http_err}")
                    if 'response' in locals():
                        st.write(response.text)  # Mostrar la respuesta completa para depuración
                except Exception as err:
                    st.error(f"Ocurrió un error: {err}")
                return "Lo siento, ocurrió un error al generar el cuento."

            # Función para generar una ilustración usando Together.xyz
            def generate_image(prompt_description):
                together_api_key = st.secrets.get('TOGETHER_API_KEY')
                if not together_api_key:
                    st.error("La clave API de Together.xyz no está configurada en los secretos de Streamlit.")
                    return None

                api_url = "https://api.together.xyz/v1/images/generations"
                headers = {
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {together_api_key}"
                }

                data = {
                    "model": "black-forest-labs/FLUX.1.1-pro",
                    "prompt": prompt_description,
                    "width": 512,
                    "height": 512,
                    "steps": 1,
                    "n": 1,
                    "response_format": "b64_json"
                }

                try:
                    response = requests.post(api_url, headers=headers, data=json.dumps(data))
                    response.raise_for_status()
                    result = response.json()
                    b64_image = result['data'][0]['b64_json']
                    image_bytes = base64.b64decode(b64_image)
                    image = Image.open(BytesIO(image_bytes))
                    return image
                except requests.exceptions.HTTPError as http_err:
                    st.error(f"Ocurrió un error HTTP al generar la imagen: {http_err}")
                    if 'response' in locals():
                        st.write(response.text)  # Mostrar la respuesta completa para depuración
                except Exception as err:
                    st.error(f"Ocurrió un error al generar la imagen: {err}")
                return None

            # Seleccionar temas para el número de cuentos
            selected_themes = random.sample(themes, num_stories) if num_stories <= len(themes) else random.choices(themes, k=num_stories)

            # Generar todos los cuentos y sus ilustraciones
            stories = []
            used_names = []
            progress_bar = st.progress(0)
            for i in range(num_stories):
                theme = selected_themes[i]
                # Obtener un nombre de personaje único
                character_name = get_unique_name(used_names)
                used_names.append(character_name)
                story = generate_corrected_story(theme, selected_age_group, character_name)
                if story.startswith("Lo siento"):
                    images = [None, None]
                else:
                    # Crear descripciones para las ilustraciones basadas en el tema y el contenido del cuento
                    image_prompt1 = f"Una ilustración colorida y atractiva para un cuento infantil sobre {theme}."
                    image_prompt2 = f"Una segunda ilustración que represente un momento clave en la historia de {character_name} en el tema de {theme}."
                    image1 = generate_image(image_prompt1)
                    image2 = generate_image(image_prompt2)
                    images = [image1, image2]
                stories.append({
                    "title": f"Capítulo {i+1}: {theme.title()}",
                    "content": story,
                    "images": images  # Almacenar una lista de imágenes
                })
                progress_bar.progress((i + 1) / num_stories)

            # Crear un documento de Word
            doc = Document()
            doc.add_heading("Cuentos Infantiles en Latín Simplificado con Ilustraciones", 0)

            for story in stories:
                doc.add_heading(story["title"], level=1)
                doc.add_paragraph(story["content"])
                if story["images"]:
                    for idx, image in enumerate(story["images"]):
                        if image:
                            # Guardar la imagen en un flujo de bytes
                            img_byte_arr = BytesIO()
                            image.save(img_byte_arr, format='PNG')
                            img_byte_arr = img_byte_arr.getvalue()
                            
                            # Agregar la imagen al documento
                            doc.add_picture(BytesIO(img_byte_arr), width=Inches(4))
                            doc.add_paragraph("")  # Espacio adicional después de la imagen
                        else:
                            doc.add_paragraph(f"No se pudo generar la ilustración {idx + 1}.")
                else:
                    doc.add_paragraph("No se pudieron generar las ilustraciones para este cuento.")

            # Guardar el documento en un flujo de BytesIO
            doc_io = BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)

            # Proporcionar el botón de descarga
            st.success("¡Cuentos e ilustraciones generados con éxito!")
            st.download_button(
                label="📄 Descargar Cuentos como Documento de Word",
                data=doc_io,
                file_name="Cuentos_Infantiles_en_Latin_Simplificado_Con_Ilustraciones.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

            # Opcional: mostrar los cuentos e ilustraciones en la aplicación
            with st.expander("Ver Cuentos e Ilustraciones Generados"):
                for story in stories:
                    st.markdown(f"### {story['title']}")
                    st.write(story["content"])
                    if story["images"]:
                        for idx, image in enumerate(story["images"]):
                            if image:
                                st.image(image, caption=f"Ilustración {idx + 1} del cuento", use_column_width=True)
                            else:
                                st.write(f"No se pudo generar la ilustración {idx + 1}.")
                    else:
                        st.write("No se pudieron generar las ilustraciones para este cuento.")

    # Pie de página
    st.markdown("---")
    st.markdown("© 2024 Generador de Cuentos Infantiles. Desarrollado con OpenRouter, Together.xyz y Streamlit.")
    ```

## **3. Explicación Detallada de las Modificaciones**

### **a. Actualización de los Prompts para Latín Simplificado**

- **Objetivo:** Instruir al modelo para que genere las historias en **latín simplificado**.
- **Implementación:** En la función `generate_corrected_story`, el prompt ahora incluye la instrucción explícita para escribir en latín simplificado.
  
  ```python
  prompt = f"""
  {guidelines}

  Crea una historia que cumpla con las siguientes características:

  - Tema: {theme}
  - Grupo de edad: {age_group}
  - Nombre del personaje principal: {character_name}
  - Texto continuo sin subdivisiones ni subtítulos.
  - Escribe la historia en **latín simplificado**, utilizando una gramática y vocabulario adecuados para niños que están aprendiendo el idioma.
  """

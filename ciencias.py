# app.py

import streamlit as st
import requests
import time
from docx import Document
from docx.shared import Inches
import base64
from io import BytesIO
import re

# Función para generar el contenido del capítulo usando la API de OpenRouter
def generate_chapter(topic):
    api_url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
    }
    prompt = (
        f"Write a chapter for a children's book about the scientific discipline specified by the user, suitable for kids aged 9 to 12. "
        f"The chapter should be both educational and engaging, presenting scientific concepts in an entertaining and informative manner. "
        f"It should aim to spark curiosity and interest in the specified scientific field while being accessible and enjoyable for the target age group.\n\n"
        f"Your chapter should blend storytelling with educational content, using creative and engaging language to explain scientific principles and concepts. "
        f"It should incorporate interactive elements, such as experiments or activities, to make learning fun and hands-on for the young readers.\n\n"
        f"Please ensure that the chapter is age-appropriate, presenting scientific ideas in a way that is understandable and captivating for children aged 9 to 12. "
        f"Consider using storytelling, illustrations, and interactive elements to create an immersive and educational experience for the young readers while maintaining a balance between instruction and entertainment.\n\n"
        f"Additionally, cover the following aspects:\n"
        f"1. Introduction to the scientific discipline ({topic}): its definition and purpose.\n"
        f"2. Scope of {topic}: what areas it covers.\n"
        f"3. Methods used in {topic}: how scientists conduct research in this field.\n"
        f"4. Achievements in {topic}: significant discoveries and advancements.\n"
        f"5. Current state of {topic}: latest trends and developments.\n"
        f"6. Include numerous examples to illustrate concepts."
    )
    data = {
        "model": "openai/gpt-4o-mini",
        "messages": [
            {
                "role": "user",
                "content": prompt
            }
        ]
    }

    response = requests.post(api_url, headers=headers, json=data)
    if response.status_code == 200:
        return response.json()['choices'][0]['message']['content']
    else:
        st.error(f"Error al generar el capítulo: {response.text}")
        return None

# Función para generar una ilustración usando la API de Together
def generate_image(topic, chapter_number, img_num):
    api_url = "https://api.together.xyz/v1/images/generations"
    headers = {
        "Authorization": f"Bearer {st.secrets['TOGETHER_API_KEY']}",
        "Content-Type": "application/json"
    }
    prompt = (
        f"Create an engaging and colorful illustration for a children's science book chapter about {topic}. "
        f"Illustration {img_num} for Chapter {chapter_number}. "
        f"The illustration should complement the storytelling and educational content, making it appealing and understandable for children aged 9 to 12."
    )
    data = {
        "model": "black-forest-labs/FLUX.1.1-pro",
        "prompt": prompt,
        "width": 1024,
        "height": 768,
        "steps": 50,  # Aumentado para mejorar la calidad de la imagen
        "n": 1,
        "response_format": "b64_json"
    }

    response = requests.post(api_url, headers=headers, json=data)
    if response.status_code == 200:
        image_b64 = response.json()['data'][0]['b64_json']
        image_bytes = base64.b64decode(image_b64)
        return image_bytes
    else:
        st.error(f"Error al generar la imagen: {response.text}")
        return None

# Función para procesar el contenido con formato Markdown y aplicarlo en el documento Word
def add_formatted_content(doc, content):
    lines = content.split('\n')
    for line in lines:
        # Encabezados de nivel 1 (equivale a #)
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        # Encabezados de nivel 2 (equivale a ##)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        # Encabezados de nivel 3 (equivale a ###)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        else:
            # Procesar texto con negrita **texto**
            # Utilizar expresiones regulares para encontrar textos entre **
            bold_pattern = re.compile(r'\*\*(.*?)\*\*')
            parts = bold_pattern.split(line)
            paragraph = doc.add_paragraph()
            for idx, part in enumerate(parts):
                if idx % 2 == 1:
                    # Texto en negrita
                    run = paragraph.add_run(part)
                    run.bold = True
                else:
                    # Texto normal
                    paragraph.add_run(part)

# Función para crear el documento de Word
def create_word_document(book_title, chapters):
    doc = Document()
    doc.add_heading(book_title, 0)

    for idx, chapter in enumerate(chapters, 1):
        # Añadir el título del capítulo como encabezado de nivel 1
        doc.add_heading(f"Chapter {idx}: {chapter['topic'].capitalize()}", level=1)
        
        # Insertar la primera ilustración al inicio del capítulo
        if len(chapter['images']) >= 1:
            doc.add_paragraph(f"Figure 1: Illustration for Chapter {idx} (Beginning)")
            image_stream = BytesIO(chapter['images'][0])
            doc.add_picture(image_stream, width=Inches(6))
        
        # Añadir el contenido del capítulo con formato
        add_formatted_content(doc, chapter['content'])

        # Insertar la segunda ilustración al final del capítulo
        if len(chapter['images']) >= 2:
            doc.add_paragraph(f"Figure 2: Illustration for Chapter {idx} (End)")
            image_stream = BytesIO(chapter['images'][1])
            doc.add_picture(image_stream, width=Inches(6))

    return doc

# Configuración de la interfaz de la aplicación
st.title("Generador de Libros de Ciencia para Niños")
st.write("Crea un libro de ciencia personalizado para niños de 9 a 12 años con capítulos e ilustraciones generadas automáticamente.")

# Entradas del usuario
book_title = st.text_input("Título del Libro", "Aventuras Científicas Asombrosas")
num_chapters = st.number_input("Número de Capítulos", min_value=1, max_value=20, value=5, step=1)
topics = st.text_area(
    "Lista de Temas (uno por línea)",
    "Astronomía\nBiología\nGeología\nBotánica\nFísica\nQuímica\nCiencias Ambientales\nBiología Marina\nTecnología\nExploración Espacial"
)

# Botón para iniciar la generación del libro
if st.button("Generar Libro"):
    topics_list = [topic.strip() for topic in topics.split('\n') if topic.strip()]
    if len(topics_list) < num_chapters:
        st.warning("No se han proporcionado suficientes temas para el número de capítulos. Algunos temas se reutilizarán.")
    
    chapters = []
    progress_bar = st.progress(0)
    status_text = st.empty()

    for i in range(num_chapters):
        topic = topics_list[i % len(topics_list)]
        status_text.text(f"Generando Capítulo {i+1} sobre {topic}...")
        content = generate_chapter(topic)
        if not content:
            st.error("Error al generar el contenido del capítulo. Se detiene la generación.")
            break

        images = []
        for img_num in range(1, 3):  # Generar dos ilustraciones por capítulo
            image = generate_image(topic, i+1, img_num)
            if image:
                images.append(image)
            time.sleep(1)  # Para evitar exceder los límites de las APIs

        chapters.append({
            "topic": topic,
            "content": content,
            "images": images
        })

        progress = (i + 1) / num_chapters
        progress_bar.progress(progress)
        time.sleep(0.5)  # Simulación de progreso

    if len(chapters) == num_chapters:
        status_text.text("Generando el documento de Word...")
        doc = create_word_document(book_title, chapters)

        # Guardar el documento en un stream de BytesIO
        doc_stream = BytesIO()
        doc.save(doc_stream)
        doc_bytes = doc_stream.getvalue()

        # Codificar el documento en base64 para la descarga
        b64 = base64.b64encode(doc_bytes).decode()

        href = f'<a href="data:application/octet-stream;base64,{b64}" download="{book_title.replace(" ", "_")}.docx">📄 Descargar Libro como Documento de Word</a>'
        st.markdown(href, unsafe_allow_html=True)
        st.success("¡Generación del libro completada exitosamente!")
    else:
        st.error("No se generaron todos los capítulos requeridos.")

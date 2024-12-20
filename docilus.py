import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import requests
import json
import base64
from PIL import Image
from io import BytesIO
from typing import List, Optional
from datetime import datetime

# Función para extraer texto de un documento Word basado en "CHAPTER" (case-sensitive)
def extract_chapters(docx_file) -> List[str]:
    doc = Document(docx_file)
    chapters = []
    current_chapter = ""
    for para in doc.paragraphs:
        text = para.text.strip()
        # Detectar si el párrafo comienza exactamente con "CHAPTER" en mayúsculas
        if text.startswith("CHAPTER"):
            if current_chapter:
                chapters.append(current_chapter)
            current_chapter = ""  # Iniciar un nuevo capítulo
            continue  # Saltar el párrafo que contiene "CHAPTER X"
        else:
            current_chapter += para.text + "\n"
    if current_chapter:
        chapters.append(current_chapter)
    return chapters

# Función para generar un prompt adecuado para adultos en el estilo de Gustave Doré
def generate_prompt(chapter: str) -> str:
    # Separar el capítulo en párrafos
    paragraphs = chapter.strip().split('\n')
    # Crear un resumen simple utilizando los dos primeros párrafos
    if len(paragraphs) >= 2:
        summary = ' '.join(paragraphs[:2])  # Usa los dos primeros párrafos como resumen
    elif paragraphs:
        summary = paragraphs[0]
    else:
        summary = chapter
    prompt = f"Ilustración en estilo Gustave Doré que represente: {summary}"
    return prompt

# Función para generar una imagen usando la API de Together
def generate_image(prompt: str, api_key: str) -> Optional[Image.Image]:
    url = "https://api.together.xyz/v1/images/generations"
    headers = {
        "Authorization": f"Bearer {api_key}",
        "Content-Type": "application/json"
    }
    
    # Preparar el prompt detallado
    full_prompt = f"Dibujo en estilo Gustave Doré, blanco y negro, detallado: {prompt}"
    
    # Obtener la fecha y hora actual en formato ISO para 'update_at'
    update_at = datetime.utcnow().isoformat() + "Z"
    
    payload = {
        "model": "black-forest-labs/FLUX.1-pro",
        "prompt": full_prompt,
        "width": 512,      # Ancho de la imagen en píxeles
        "height": 512,     # Alto de la imagen en píxeles
        "steps": 28,       # Número de pasos para la generación
        "n": 1,            # Número de imágenes a generar
        "response_format": "b64_json",
        "update_at": update_at
    }
    
    try:
        response = requests.post(url, headers=headers, data=json.dumps(payload))
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        st.error(f"Error al generar la imagen: {e}")
        return None
    
    data = response.json()
    image_data = data.get("data", [])[0].get("b64_json", "")
    
    if not image_data:
        st.error("No se recibió imagen en la respuesta.")
        return None
    
    try:
        # Decodificar la imagen de base64
        image_bytes = base64.b64decode(image_data)
        image = Image.open(BytesIO(image_bytes))
        
        # Verificar dimensiones de la imagen
        if image.size != (512, 512):
            st.warning(f"La imagen generada tiene un tamaño de {image.size}, redimensionando a 512x512.")
            image = image.resize((512, 512))
        
        return image
    except Exception as e:
        st.error(f"Error al procesar la imagen: {e}")
        return None

# Función para crear un nuevo documento Word con capítulos e ilustraciones
def create_word_document(chapters: List[str], images: List[Image.Image]) -> BytesIO:
    new_doc = Document()
    
    # Configurar el tamaño de página y márgenes
    section = new_doc.sections[0]
    section.page_width = Inches(7.5)
    section.page_height = Inches(7.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    
    # Definir estilos para los párrafos
    style = new_doc.styles['Normal']
    font = style.font
    font.name = 'Alegreya'  # Cambiado a 'Alegreya'
    font.size = Pt(14)
    
    # Configurar interlineado y espacio después
    paragraph_format = style.paragraph_format
    paragraph_format.line_spacing = 1.15
    paragraph_format.space_after = Pt(10)
    paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    
    for idx, (chapter, image) in enumerate(zip(chapters, images), 1):
        # Añadir el capítulo como un encabezado (Heading 1)
        new_doc.add_heading(f"Capítulo {idx}", level=1)
        
        # Añadir el texto del capítulo, separando en párrafos
        paragraphs = chapter.strip().split('\n')
        for para_text in paragraphs:
            if para_text.strip():  # Evitar párrafos vacíos
                para = new_doc.add_paragraph(para_text.strip())
                para.style = 'Normal'
        
        # Añadir la ilustración
        if image:
            # Guardar la imagen en un buffer BytesIO
            img_buffer = BytesIO()
            image.save(img_buffer, format='PNG')
            img_buffer.seek(0)
            
            # Añadir la imagen al documento con tamaño 5.5 x 5.5 pulgadas
            new_doc.add_picture(img_buffer, width=Inches(5.5), height=Inches(5.5))
            
            # Añadir un salto de párrafo después de la imagen
            new_doc.add_paragraph()
    
    # Guardar el documento en un buffer BytesIO
    doc_buffer = BytesIO()
    new_doc.save(doc_buffer)
    doc_buffer.seek(0)
    
    return doc_buffer

# Interfaz de Streamlit
def main():
    st.set_page_config(page_title="Generador de Ilustraciones para Fábulas", layout="wide")
    st.title("Generador de Ilustraciones para Fábulas")
    
    uploaded_file = st.file_uploader("Selecciona un archivo Word", type=["docx"])
    
    # Verificar si la clave API está disponible
    if "TOGETHER_API_KEY" not in st.secrets:
        st.error("Falta la clave API de Together. Por favor, agrega `TOGETHER_API_KEY` en los secretos de Streamlit.")
        st.stop()
    
    together_api_key = st.secrets["TOGETHER_API_KEY"]
    
    if uploaded_file:
        if st.button("Generar Documento con Ilustraciones"):
            with st.spinner("Procesando el documento..."):
                # Extraer capítulos del documento subido
                chapters = extract_chapters(uploaded_file)
            
            st.success(f"Se han extraído {len(chapters)} capítulo(s).")
            
            images = []
            for idx, chapter in enumerate(chapters, 1):
                prompt = generate_prompt(chapter)
                
                with st.spinner(f"Generando imagen para el Capítulo {idx}..."):
                    image = generate_image(prompt, together_api_key)
                    images.append(image)
            
            # Crear el nuevo documento Word con capítulos e ilustraciones
            with st.spinner("Creando el nuevo documento Word..."):
                doc_buffer = create_word_document(chapters, images)
            
            st.success("Documento Word generado exitosamente.")
            
            # Ofrecer el documento para descarga
            st.download_button(
                label="Descargar Documento con Ilustraciones",
                data=doc_buffer,
                file_name="fábula_con_ilustraciones.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

if __name__ == "__main__":
    main()

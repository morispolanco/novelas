import streamlit as st
import requests
import json
import time
from docx import Document
from io import BytesIO
import re
from requests.adapters import HTTPAdapter
from urllib3.util.retry import Retry
from docx.shared import Inches, Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import logging
from time import sleep

# Configurar el logging para Streamlit
logging.basicConfig(format='%(asctime)s - %(levelname)s - %(message)s', level=logging.DEBUG)  # Nivel DEBUG para más detalles

# Configuración de la página
st.set_page_config(
    page_title="Analizador de Novelas de Suspenso Político",
    layout="wide",
    initial_sidebar_state="expanded",
)

# Título de la aplicación
st.title("Analizador de Novelas de Suspenso Político")
st.write("""
Esta aplicación analiza una novela en el género de thriller político.
Sube tu novela en formato `.docx` o `.txt` y recibe un informe detallado de errores y mejoras, además de una calificación global.
""")

# Barra lateral para opciones de usuario
st.sidebar.header("Configuración de Análisis")
file_upload = st.sidebar.file_uploader("Sube tu novela (.docx o .txt):", type=["docx", "txt"])

# Inicializar el estado de la aplicación
if 'etapa' not in st.session_state:
    st.session_state.etapa = "inicio"  # etapas: inicio, analisis, completado

if 'novela' not in st.session_state:
    st.session_state.novela = ""
if 'informe_global' not in st.session_state:
    st.session_state.informe_global = ""

# Función para extraer JSON de una respuesta mixta
def extraer_json(texto):
    try:
        # Encuentra el primer bloque que comienza con { y termina con }
        match = re.search(r'\{.*\}', texto, re.DOTALL)
        if match:
            json_str = match.group(0)
            return json.loads(json_str)
        else:
            return None
    except json.JSONDecodeError:
        return None

# Función para llamar a la API de OpenRouter con reintentos y parámetros ajustables
def call_openrouter_api(prompt, max_tokens=3000, temperature=0.5, top_p=0.9, repetition_penalty=1.2):
    api_url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "openai/gpt-4o-mini",  # Usando el modelo especificado por el usuario
        "messages": [{"role": "user", "content": prompt}],
        "temperature": temperature,
        "max_tokens": max_tokens,
        "top_p": top_p,
        "repetition_penalty": repetition_penalty,
        "stream": False
        # "stop": ["<|eot_id|>"],  # Eliminado para evitar posibles interrupciones incorrectas
    }
    
    session = requests.Session()
    retries = Retry(total=5, backoff_factor=1, status_forcelist=[502, 503, 504])
    session.mount('https://', HTTPAdapter(max_retries=retries))
    
    response_text = ""
    for attempt in range(5):
        try:
            response = session.post(api_url, headers=headers, data=json.dumps(payload))
            response.raise_for_status()
            response_text = response.text  # Obtener la respuesta como texto
            logging.debug("Respuesta de la API: %s", response_text)  # Registrar la respuesta completa
            
            # Mostrar detalles de la respuesta para depuración
            logging.debug("Código de estado HTTP: %s", response.status_code)
            logging.debug("Encabezados de la respuesta: %s", response.headers)
            
            if 'application/json' in response.headers.get('Content-Type', ''):
                response_json = response.json()
            else:
                response_json = extraer_json(response_text)
                if not response_json:
                    st.error("La respuesta de la API no contiene un JSON válido.")
                    st.write("**Respuesta completa de la API:**", response_text)
                    return None
            
            if 'choices' in response_json and len(response_json['choices']) > 0:
                contenido = response_json['choices'][0]['message']['content']
                logging.debug("Contenido de la respuesta de la API: %s", contenido)
                return contenido
            else:
                st.error("La respuesta de la API no contiene 'choices'.")
                st.write("**Respuesta completa de la API:**", response_json)
                return None
        except requests.exceptions.RequestException as e:
            logging.error(f"Intento {attempt + 1}: Error en la llamada a la API: {e}")
            st.warning(f"Intento {attempt + 1}: Error en la llamada a la API. Reintentando...")
            sleep(2 ** attempt)
        except json.JSONDecodeError as e:
            logging.error("Error al decodificar la respuesta de la API: %s", e)
            st.error(f"Error al decodificar la respuesta de la API: {e}")
            st.write("**Respuesta de la API (texto):**", response_text)
            return None
    
    st.error("No se pudo obtener una respuesta válida de la API después de varios intentos.")
    return None

# Función para leer el contenido del archivo subido
def leer_archivo(file):
    try:
        if file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(file)
            texto = "\n".join([para.text for para in doc.paragraphs])
        elif file.type == "text/plain":
            texto = file.read().decode("utf-8")
        else:
            st.error("Formato de archivo no soportado.")
            return None
        return texto
    except Exception as e:
        st.error(f"Error al leer el archivo: {e}")
        logging.error("Error al leer el archivo: %s", e)
        return None

# Función para dividir la novela si es muy larga
def dividir_novela(texto, max_length=3000):
    # Dividir por párrafos para mantener la coherencia
    parrafos = texto.split('\n')
    secciones = []
    seccion_actual = ""
    
    for parrafo in parrafos:
        if len(seccion_actual) + len(parrafo) + 1 <= max_length:
            seccion_actual += parrafo + '\n'
        else:
            secciones.append(seccion_actual.strip())
            seccion_actual = parrafo + '\n'
    
    if seccion_actual:
        secciones.append(seccion_actual.strip())
    
    return secciones

# Función para analizar la novela completa
def analizar_novela(texto, progress_bar=None, progress_text=None):
    secciones = dividir_novela(texto)
    total_secciones = len(secciones)
    analisis_completo = {
        "calificacion": [],
        "errores": [],
        "recomendaciones": []
    }
    
    for idx, seccion in enumerate(secciones):
        prompt = f"""
Por favor, analiza la siguiente sección de una novela de suspenso político. Identifica errores gramaticales, de coherencia, desarrollo de personajes, ritmo y cualquier otro aspecto que pueda mejorar la calidad de la novela. Proporciona recomendaciones claras y específicas para cada área de mejora y asigna una calificación de 1 a 10 puntos basada en la calidad general de la sección.

Responde únicamente con un JSON válido que contenga las siguientes claves: "calificacion", "errores", "recomendaciones". No incluyas ningún texto adicional fuera del JSON.

### Sección {idx+1}:
{seccion}

### Informe de Análisis:
"""
        analisis = call_openrouter_api(prompt)
        if analisis:
            try:
                analisis_json = json.loads(analisis)
                calificacion = analisis_json.get('calificacion', 0)
                errores = analisis_json.get('errores', 'No se identificaron errores.')
                recomendaciones = analisis_json.get('recomendaciones', 'No se proporcionaron recomendaciones.')
                
                analisis_completo["calificacion"].append(float(calificacion))
                analisis_completo["errores"].append(errores)
                analisis_completo["recomendaciones"].append(recomendaciones)
            except json.JSONDecodeError as e:
                st.error("Error al decodificar la respuesta de la API.")
                st.write("**Sección:**", idx+1)
                st.write("**Respuesta de la API:**", analisis)
                logging.error("Error al decodificar la respuesta JSON: %s", e)
                return None
        else:
            st.error("No se recibió análisis para una sección.")
            return None
        
        # Actualizar la barra de progreso y el texto
        if progress_bar and progress_text:
            progreso_actual = (idx + 1) / total_secciones
            progress_bar.progress(progreso_actual)
            progress_text.text(f"Procesando sección {idx + 1} de {total_secciones} ({int(progreso_actual * 100)}%)")
    
    # Calcular la calificación promedio global
    if analisis_completo["calificacion"]:
        calificacion_promedio = sum(analisis_completo["calificacion"]) / len(analisis_completo["calificacion"])
    else:
        calificacion_promedio = "No disponible"
    
    # Combinar errores y recomendaciones para el análisis global
    errores_completos = "\n".join([f"**Sección {i+1}:**\n{error}" for i, error in enumerate(analisis_completo["errores"])])
    recomendaciones_completas = "\n".join([f"**Sección {i+1}:**\n{recomendacion}" for i, recomendacion in enumerate(analisis_completo["recomendaciones"])])
    
    informe_global = {
        "calificacion": round(calificacion_promedio, 2) if isinstance(calificacion_promedio, float) else calificacion_promedio,
        "errores": errores_completos,
        "recomendaciones": recomendaciones_completas
    }
    
    return json.dumps(informe_global, ensure_ascii=False), None  # Retorna None para el informe de escenas

# Función para generar el informe completo
def generar_informe(informe_global, _):
    try:
        analisis_global_json = json.loads(informe_global)
        calificacion = analisis_global_json.get('calificacion', 'No disponible')
        errores = analisis_global_json.get('errores', 'No se identificaron errores.')
        recomendaciones = analisis_global_json.get('recomendaciones', 'No se proporcionaron recomendaciones.')
        
        informe = f"# Informe de Análisis de la Novela\n\n"
        informe += f"## **Calificación General:** {calificacion} / 10\n\n"
        informe += f"## **Errores Identificados:**\n{errores}\n\n"
        informe += f"## **Recomendaciones para Mejoras:**\n{recomendaciones}\n\n"
        return informe
    except json.JSONDecodeError:
        informe = f"# Informe de Análisis de la Novela\n\n"
        informe += f"**Calificación General:** No disponible\n\n"
        informe += f"**Errores y Recomendaciones:**\n{informe_global}\n\n"
        return informe

# Función para exportar el informe a Word
def exportar_informe_word(informe):
    document = Document()
    section = document.sections[0]
    section.page_width = Inches(6)
    section.page_height = Inches(9)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)

    style = document.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    titulo_paragraph = document.add_heading("Informe de Análisis de la Novela", level=0)
    titulo_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    document.add_page_break()

    for line in informe.split('\n'):
        if line.startswith("# "):
            document.add_heading(line.replace("# ", ""), level=1)
        elif line.startswith("## "):
            document.add_heading(line.replace("## ", ""), level=2)
        elif line.startswith("**"):
            match = re.match(r'\*\*(.*?)\*\*:\s*(.*)', line)
            if match:
                term, desc = match.groups()
                p = document.add_paragraph()
                p.add_run(f"{term}: ").bold = True
                p.add_run(desc)
        else:
            if line.strip() != "":
                document.add_paragraph(line)

    buffer = BytesIO()
    document.save(buffer)
    buffer.seek(0)
    return buffer

# Función para probar la API con un prompt simplificado
def probar_api():
    prompt = """
Por favor, responde únicamente con un JSON que contenga las siguientes claves: "mensaje".

### Solicitud:
Hola, ¿cómo estás?

### Respuesta:
"""
    respuesta = call_openrouter_api(prompt)
    if respuesta:
        try:
            respuesta_json = json.loads(respuesta)
            st.write("**Respuesta de Prueba:**", respuesta_json)
        except json.JSONDecodeError as e:
            st.error("Error al decodificar la respuesta de prueba de la API.")
            st.write("**Respuesta de la API (texto):**", respuesta)
            logging.error("Error al decodificar la respuesta JSON de prueba: %s", e)

# Interfaz de usuario para cargar y analizar
def mostrar_inicio():
    st.header("Carga y Análisis de la Novela")
    with st.form(key='form_carga'):
        archivo = st.file_uploader("Sube tu novela (.docx o .txt):", type=["docx", "txt"])
        submit_btn = st.form_submit_button(label='Enviar')
    
    if submit_btn:
        if archivo is not None:
            texto = leer_archivo(archivo)
            if texto:
                st.session_state.novela = texto
                st.session_state.etapa = "analisis"
                st.success("Archivo cargado y listo para análisis.")
        else:
            st.error("Por favor, carga un archivo antes de hacer clic en 'Enviar'.")
    
    st.markdown("---")
    st.header("Prueba de la API")
    probar_btn = st.button("Probar API con Prompt Simplificado")
    
    if probar_btn:
        with st.spinner("Probando la API..."):
            probar_api()

# Interfaz de usuario para el análisis
def mostrar_analisis():
    st.header("Análisis en Proceso")
    novela = st.session_state.novela
    st.subheader("Extracto de la Novela:")
    st.text_area("Contenido de la Novela:", novela[:1000] + "..." if len(novela) > 1000 else novela, height=200)

    iniciar_btn = st.button("Iniciar Análisis")

    if iniciar_btn:
        with st.spinner("Analizando la novela..."):
            progreso = st.progress(0)
            progreso_text = st.empty()
            analisis_global, _ = analizar_novela(novela, progress_bar=progreso, progress_text=progreso_text)
            if analisis_global:
                informe = generar_informe(analisis_global, None)
                st.session_state.informe_global = informe
                st.session_state.etapa = "completado"
                st.success("Análisis completado exitosamente.")
            else:
                st.error("No se pudo generar el análisis. Por favor, intenta nuevamente.")

# Interfaz de usuario para mostrar el informe y descargar
def mostrar_completado():
    st.header("Informe de Análisis Generado")
    informe_global = st.session_state.informe_global

    st.subheader("Informe Global:")
    st.text_area("Informe Global:", informe_global, height=600)

    # Exportar a Word
    doc_buffer = exportar_informe_word(informe_global)
    st.download_button(
        label="Descargar Informe en Word",
        data=doc_buffer,
        file_name=f"informe_analisis_novela_{int(time.time())}.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

# Interfaz de usuario principal
if st.session_state.etapa == "inicio":
    mostrar_inicio()

elif st.session_state.etapa == "analisis":
    mostrar_analisis()

elif st.session_state.etapa == "completado":
    mostrar_completado()

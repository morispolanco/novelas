import streamlit as st
import requests
import time
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from io import BytesIO
import json

# Configuración de la página
st.set_page_config(page_title="Generador de Novelas de Thriller Político", layout="wide")

# Función para llamar a la API de OpenRouter
def call_openrouter_api(messages):
    api_url = "https://openrouter.ai/api/v1/chat/completions"
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {st.secrets['OPENROUTER_API_KEY']}"
    }
    data = {
        "model": "openai/gpt-4o-mini",  # Verifica el nombre correcto del modelo
        "messages": messages
    }
    try:
        response = requests.post(api_url, headers=headers, json=data)
        response.raise_for_status()
    except requests.exceptions.HTTPError as http_err:
        st.error(f"HTTP error occurred: {http_err}")
        st.stop()
    except Exception as err:
        st.error(f"Other error occurred: {err}")
        st.stop()
    
    try:
        response_json = response.json()
    except json.JSONDecodeError:
        st.error("La respuesta de la API no es un JSON válido.")
        st.write("Respuesta completa de la API:", response.text)
        st.stop()
    
    try:
        content = response_json["choices"][0]["message"]["content"].strip()
    except (KeyError, IndexError) as e:
        st.error(f"Error al parsear la respuesta de la API: {e}")
        st.write("Respuesta completa de la API:", response_json)
        st.stop()
    
    return content

# Función para planificar la novela
def plan_novel(theme):
    prompt = [
        {"role": "user", "content": f"Planifica una novela de thriller político basada en el siguiente tema: {theme}. Genera la trama principal y las subtramas."}
    ]
    return call_openrouter_api(prompt)

# Función para distribuir la trama en capítulos y secciones
def distribute_plot(plot):
    prompt = [
        {"role": "user", "content": (
            "Distribuye la siguiente trama en 12 capítulos, cada uno con 5 secciones. "
            "Proporciona una lista estructurada con capítulos y secciones en formato JSON, "
            "asegurándote de que el JSON esté bien formado y sea válido.\n\n"
            f"Trama: {plot}"
        )}
    ]
    return call_openrouter_api(prompt)

# Función para describir cada sección
def describe_sections(distribution):
    prompt = [
        {"role": "user", "content": (
            "Describe detalladamente qué sucede en cada una de las siguientes 60 secciones para asegurar coherencia y consistencia en la trama. "
            "Proporciona las descripciones en formato JSON, asegurándote de que el JSON esté bien formado y sea válido.\n\n"
            f"Distribución: {distribution}"
        )}
    ]
    return call_openrouter_api(prompt)

# Función para escribir una sección
def write_section(section_description):
    prompt = [
        {"role": "user", "content": (
            "Escribe una sección de aproximadamente 1000 palabras con las siguientes características de un thriller político: mucha acción, ritmo rápido, "
            "descripciones vívidas, finales de escena con ganchos para mantener la atención del lector. Evita descripciones excesivas y frases cliché. "
            "Utiliza raya para los diálogos.\n\n"
            f"Descripción de la sección: {section_description}"
        )}
    ]
    return call_openrouter_api(prompt)

# Función mejorada para crear el documento Word
def create_word_document(chapters, sections_content):
    # Crear un nuevo documento
    doc = Document()
    
    # Definir estilos personalizados
    styles = doc.styles
    
    # Estilo para el título principal
    if 'TitleStyle' not in styles:
        title_style = styles.add_style('TitleStyle', WD_STYLE_TYPE.PARAGRAPH)
        title_style.font.name = 'Arial'
        title_style.font.size = Pt(24)
        title_style.font.bold = True
        title_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Estilo para los capítulos
    if 'ChapterStyle' not in styles:
        chapter_style = styles.add_style('ChapterStyle', WD_STYLE_TYPE.PARAGRAPH)
        chapter_style.font.name = 'Arial'
        chapter_style.font.size = Pt(18)
        chapter_style.font.bold = True
        chapter_style.paragraph_format.space_before = Pt(12)
        chapter_style.paragraph_format.space_after = Pt(6)
    
    # Estilo para las secciones
    if 'SectionStyle' not in styles:
        section_style = styles.add_style('SectionStyle', WD_STYLE_TYPE.PARAGRAPH)
        section_style.font.name = 'Arial'
        section_style.font.size = Pt(14)
        section_style.font.bold = True
        section_style.paragraph_format.space_before = Pt(6)
        section_style.paragraph_format.space_after = Pt(3)
    
    # Estilo para el contenido de las secciones
    if 'ContentStyle' not in styles:
        content_style = styles.add_style('ContentStyle', WD_STYLE_TYPE.PARAGRAPH)
        content_style.font.name = 'Calibri'
        content_style.font.size = Pt(12)
        content_style.paragraph_format.space_before = Pt(0)
        content_style.paragraph_format.space_after = Pt(6)
        content_style.paragraph_format.line_spacing = 1.15
    
    # Añadir título principal
    title = doc.add_paragraph("Novela de Thriller Político", style='TitleStyle')
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    doc.add_paragraph()  # Línea en blanco
    
    # Iterar sobre los capítulos
    for chapter_num, chapter in enumerate(chapters, start=1):
        # Añadir título del capítulo
        chapter_title = doc.add_paragraph(f"Capítulo {chapter_num}: {chapter['title']}", style='ChapterStyle')
        
        # Iterar sobre las secciones del capítulo
        for section_num, section in enumerate(chapter['sections'], start=1):
            # Añadir título de la sección
            section_title = doc.add_paragraph(f"Sección {chapter_num}.{section_num}: {section['title']}", style='SectionStyle')
            
            # Añadir contenido de la sección
            section_key = f"{chapter_num}.{section_num}"
            content_text = sections_content.get(section_key, "Contenido no disponible.")
            content_paragraph = doc.add_paragraph(content_text, style='ContentStyle')
    
    # Agregar número de página en el pie de página
    section = doc.sections[-1]
    footer = section.footer
    footer_paragraph = footer.paragraphs[0]
    footer_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    footer_paragraph.add_run("Página ").italic = True
    footer_paragraph.add_run().add_field('PAGE')
    
    # Guardar el documento en un objeto BytesIO
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)
    return byte_io

# Título de la aplicación
st.title("Generador de Novelas de Thriller Político")

# Entrada para el tema de la novela
theme = st.text_input("Ingresa el tema de tu novela de thriller político:", "")

# Inicializar variables de estado
if 'plan' not in st.session_state:
    st.session_state.plan = None
if 'distribution' not in st.session_state:
    st.session_state.distribution = None
if 'sections_description' not in st.session_state:
    st.session_state.sections_description = None
if 'sections_content' not in st.session_state:
    st.session_state.sections_content = {}
if 'chapters' not in st.session_state:
    st.session_state.chapters = []
if 'step' not in st.session_state:
    st.session_state.step = 0
if 'completed' not in st.session_state:
    st.session_state.completed = False

# Función para reiniciar el proceso
def reset_process():
    st.session_state.plan = None
    st.session_state.distribution = None
    st.session_state.sections_description = None
    st.session_state.sections_content = {}
    st.session_state.chapters = []
    st.session_state.step = 0
    st.session_state.completed = False

# Botón para reiniciar
if st.button("Reiniciar Proceso"):
    reset_process()
    st.success("Proceso reiniciado.")

# Botón para iniciar el proceso
if st.button("Generar Novela"):
    if not theme.strip():
        st.error("Por favor, ingresa un tema para la novela.")
    else:
        with st.spinner("Planificando la trama de la novela..."):
            plot = plan_novel(theme)
            st.session_state.plan = plot
            st.session_state.step = 1

# Mostrar el plan para aprobación si está disponible
if st.session_state.step == 1 and st.session_state.plan:
    st.subheader("Plan de la Novela")
    st.write(st.session_state.plan)
    st.write("¿Deseas aprobar este plan y continuar con la generación de la novela?")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Aprobar y Continuar"):
            st.session_state.step = 2
    with col2:
        if st.button("Cancelar"):
            reset_process()
            st.info("Proceso cancelado.")

# Continuar con la distribución de la trama si el plan fue aprobado
if st.session_state.step == 2:
    # Barra de progreso
    progress_bar = st.progress(0)
    status_text = st.empty()

    # Paso 2: Distribuir la trama en capítulos y secciones
    status_text.text("Distribuyendo la trama en capítulos y secciones...")
    distribution = distribute_plot(st.session_state.plan)
    st.session_state.distribution = distribution
    progress_bar.progress(20)
    time.sleep(1)

    # Paso 3: Describir cada sección
    status_text.text("Describiendo cada sección...")
    sections_description = describe_sections(st.session_state.distribution)
    st.session_state.sections_description = sections_description
    progress_bar.progress(40)
    time.sleep(1)

    # Procesar la distribución y las descripciones
    try:
        distribution_data = json.loads(st.session_state.distribution)
        sections_description_data = json.loads(st.session_state.sections_description)
    except json.JSONDecodeError:
        st.error("Error al procesar la respuesta de la API. Asegúrate de que el formato sea JSON válido.")
        st.write("Distribución recibida:", st.session_state.distribution)
        st.write("Descripción de secciones recibida:", st.session_state.sections_description)
        st.stop()

    # Crear una estructura de capítulos y secciones
    st.session_state.chapters = []
    for chapter in distribution_data.get("capítulos", []):
        chapter_dict = {
            "title": chapter.get("título", f"Capítulo {len(st.session_state.chapters)+1}"),
            "sections": [{"title": sec.get("título", f"Sección {i+1}")} for i, sec in enumerate(chapter.get("secciones", []))]
        }
        st.session_state.chapters.append(chapter_dict)

    # Asignar descripciones a cada sección
    for idx, section_desc in enumerate(sections_description_data.get("secciones", []), start=1):
        chapter_num = (idx - 1) // 5 + 1
        section_num = (idx - 1) % 5 + 1
        section_key = f"{chapter_num}.{section_num}"
        st.session_state.sections_content[section_key] = section_desc.get("descripción", "")

    progress_bar.progress(50)
    status_text.text("Descripción de secciones completada.")
    time.sleep(1)

    # Mostrar las descripciones de las escenas para aprobación
    st.subheader("Descripciones de las Escenas")
    st.write("A continuación se presentan las descripciones generadas para cada sección. ¿Deseas aprobarlas y continuar con la generación de las escenas?")
    
    # Mostrar las descripciones en un formato legible
    for chapter in st.session_state.chapters:
        st.markdown(f"### {chapter['title']}")
        for section in chapter['sections']:
            section_key = f"{st.session_state.chapters.index(chapter)+1}.{chapter['sections'].index(section)+1}"
            st.markdown(f"#### {section_key}: {section['title']}")
            st.write(st.session_state.sections_content.get(section_key, "Descripción no disponible."))
    
    # Botones de aprobación
    st.write("¿Deseas aprobar estas descripciones y proceder a generar las escenas?")
    col1, col2 = st.columns(2)
    with col1:
        if st.button("Aprobar y Continuar"):
            st.session_state.step = 3
    with col2:
        if st.button("Cancelar"):
            reset_process()
            st.info("Proceso cancelado.")

# Continuar con la generación de escenas si las descripciones fueron aprobadas
if st.session_state.step == 3:
    # Barra de progreso
    progress_bar = st.progress(0)
    status_text = st.empty()

    # Paso 4: Escribir cada sección
    status_text.text("Escribiendo cada sección de la novela...")
    total_sections = 12 * 5  # 12 capítulos, 5 secciones cada uno
    for idx in range(1, total_sections + 1):
        chapter_num = (idx - 1) // 5 + 1
        section_num = (idx - 1) % 5 + 1
        section_key = f"{chapter_num}.{section_num}"
        description = st.session_state.sections_content.get(section_key, "")
        if not description:
            st.error(f"Descripción faltante para la sección {section_key}.")
            break
        content = write_section(description)
        st.session_state.sections_content[section_key] = content
        progress = 50 + (50 * idx) // total_sections
        progress_bar.progress(progress)
        status_text.text(f"Escribiendo sección {idx} de {total_sections}...")
        time.sleep(0.5)  # Pausa para no sobrecargar la API

    progress_bar.progress(100)
    status_text.text("Generación de la novela completada.")
    st.session_state.step = 4

    # Paso 5: Crear y ofrecer el documento Word
    with st.spinner("Creando el documento Word..."):
        word_doc = create_word_document(st.session_state.chapters, st.session_state.sections_content)
    st.success("¡Novela generada exitosamente!")
    st.download_button(
        label="Descargar Novela en Word",
        data=word_doc,
        file_name="novela_thriller_politico.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )

    # Finalizar el proceso
    st.session_state.completed = True

# Mostrar progreso si el proceso está en curso
if st.session_state.step == 4 and st.session_state.completed:
    st.balloons()
